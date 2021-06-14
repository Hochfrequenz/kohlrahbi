import re
from typing import List

import pandas as pd
from docx.text.paragraph import Paragraph

from check_row_type import RowType


def parse_paragraph_in_edifact_struktur_column_to_dataframe(
    paragraph: Paragraph,
    dataframe: pd.DataFrame,
    row_index: int,
    edifact_struktur_cell_left_indent_position: int,
):

    splitted_text_at_tabs = paragraph.text.split("\t")
    tab_count = paragraph.text.count("\t")

    # Check if the line starts on the far left
    if paragraph.paragraph_format.left_indent != edifact_struktur_cell_left_indent_position:

        if tab_count == 2:
            dataframe.at[row_index, "Segment Gruppe"] = splitted_text_at_tabs[0]
            dataframe.at[row_index, "Segment"] = splitted_text_at_tabs[1]
            dataframe.at[row_index, "Datenelement"] = splitted_text_at_tabs[2]
        elif tab_count == 1:
            dataframe.at[row_index, "Segment Gruppe"] = splitted_text_at_tabs[0]
            dataframe.at[row_index, "Segment"] = splitted_text_at_tabs[1]
        elif tab_count == 0 and not paragraph.text == "":
            if paragraph.runs[0].bold:
                # Segmentgruppe: SG8
                dataframe.at[row_index, "Segment Gruppe"] = splitted_text_at_tabs[0]
            else:
                # Segmentname: Referenzen auf die ID der\nTranche
                dataframe.at[row_index, "Segment Gruppe"] += splitted_text_at_tabs[0]

    # Now the text should start in middle of the EDIFACT Struktur column
    else:

        if tab_count == 1:
            # Example: "UNH\t0062"
            dataframe.at[row_index, "Segment"] = splitted_text_at_tabs[0]
            dataframe.at[row_index, "Datenelement"] = splitted_text_at_tabs[1]

        elif tab_count == 0:
            # Example: "UNH"
            dataframe.at[row_index, "Segment"] = splitted_text_at_tabs[0]


def parse_paragraph_in_middle_column_to_dataframe(
    paragraph: Paragraph,
    dataframe: pd.DataFrame,
    row_index: int,
    left_indent_position: int,
    tabstop_positions: List,
):

    splitted_text_at_tabs = paragraph.text.split("\t")

    # Qualifier / Code
    # left_indent_position is characteristic for Datenelemente
    if paragraph.paragraph_format.left_indent == left_indent_position:
        dataframe.at[row_index, "Codes und Qualifier"] += splitted_text_at_tabs.pop(0)
        column_indezes = list(range(4, 4 + len(tabstop_positions)))

    else:
        if splitted_text_at_tabs[0] == "":
            tabstop_positions = tabstop_positions[1:]
            del splitted_text_at_tabs[0]

        column_indezes = list(range(5, 5 + len(tabstop_positions)))

    tab_stops = paragraph.paragraph_format.tab_stops._pPr.tabs

    if tab_stops is not None:
        for tabstop in tab_stops:
            for tabstop_position, column_index in zip(tabstop_positions, column_indezes):
                if tabstop.pos == tabstop_position:
                    dataframe.iat[row_index, column_index] += splitted_text_at_tabs.pop(0)
    elif tab_stops is None and splitted_text_at_tabs:
        # in splitted_text_at_tabs list must be an entry
        dataframe.at[row_index, "Beschreibung"] += splitted_text_at_tabs.pop(0)
    elif tab_stops is None:
        pass
    # Could not figure out a scenario where this error could be raised.
    # else:
    #     raise NotImplementedError(f"Could not parse paragraphe in middle cell with {paragraph.text}")


def parse_bedingung_cell(bedingung_cell, dataframe, row_index):

    bedingung = bedingung_cell.text.replace("\n", " ")
    # TODO regex condition auslagern
    matches = re.findall(r"\[\d*\]", bedingung)
    for match in matches[1:]:
        index = bedingung.find(match)
        bedingung = bedingung[:index] + "\n" + bedingung[index:]

    dataframe.at[row_index, "Bedingung"] += bedingung
    pass


def write_segment_name_to_dataframe(
    dataframe,
    row_index,
    edifact_struktur_cell,
    edifact_struktur_cell_left_indent_position,
    middle_cell,
    middle_cell_left_indent_position: int,
    tabstop_positions: List,
    bedingung_cell,
):

    # EDIFACT STRUKTUR COLUMN
    parse_paragraph_in_edifact_struktur_column_to_dataframe(
        paragraph=edifact_struktur_cell.paragraphs[0],
        dataframe=dataframe,
        row_index=row_index,
        edifact_struktur_cell_left_indent_position=edifact_struktur_cell_left_indent_position,
    )

    # MIDDLE COLUMN
    # I do not expect to a multiline Segementgruppe,
    # but just in case we loop through all paragraphs
    for paragraph in middle_cell.paragraphs:
        parse_paragraph_in_middle_column_to_dataframe(
            paragraph=paragraph,
            dataframe=dataframe,
            row_index=row_index,
            left_indent_position=middle_cell_left_indent_position,
            tabstop_positions=tabstop_positions,
        )

    # BEDINGUNG COLUMN
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=dataframe, row_index=row_index)


def write_segmentgruppe_to_dataframe(
    dataframe,
    row_index,
    edifact_struktur_cell,
    edifact_struktur_cell_left_indent_position,
    middle_cell,
    middle_cell_left_indent_position: int,
    tabstop_positions: List,
    bedingung_cell,
):
    # EDIFACT STRUKTUR COLUMN
    parse_paragraph_in_edifact_struktur_column_to_dataframe(
        paragraph=edifact_struktur_cell.paragraphs[0],
        dataframe=dataframe,
        row_index=row_index,
        edifact_struktur_cell_left_indent_position=edifact_struktur_cell_left_indent_position,
    )

    # MIDDLE COLUMN
    # I do not expect to a multiline Segementgruppe,
    # but just in case we loop through all paragraphs
    for paragraph in middle_cell.paragraphs:
        parse_paragraph_in_middle_column_to_dataframe(
            paragraph=paragraph,
            dataframe=dataframe,
            row_index=row_index,
            left_indent_position=middle_cell_left_indent_position,
            tabstop_positions=tabstop_positions,
        )

    # BEDINGUNG COLUMN
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=dataframe, row_index=row_index)


def write_segment_to_dataframe(
    dataframe,
    row_index,
    edifact_struktur_cell,
    edifact_struktur_cell_left_indent_position,
    middle_cell,
    middle_cell_left_indent_position: int,
    tabstop_positions: List,
    bedingung_cell,
):

    # EDIFACT STRUKTUR COLUMN
    parse_paragraph_in_edifact_struktur_column_to_dataframe(
        paragraph=edifact_struktur_cell.paragraphs[0],
        dataframe=dataframe,
        row_index=row_index,
        edifact_struktur_cell_left_indent_position=edifact_struktur_cell_left_indent_position,
    )

    # MIDDLE COLUMN
    for paragraph in middle_cell.paragraphs:
        parse_paragraph_in_middle_column_to_dataframe(
            paragraph=paragraph,
            dataframe=dataframe,
            row_index=row_index,
            left_indent_position=middle_cell_left_indent_position,
            tabstop_positions=tabstop_positions,
        )

    # BEDINGUNG COLUMN
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=dataframe, row_index=row_index)


def count_matching(condition, seq):
    """Returns the amount of items in seq that return true from condition"""
    return sum(condition(item) for item in seq)


def code_condition(paragraph):
    try:
        tabstop_positions = [tab_position.pos for tab_position in paragraph.paragraph_format.tab_stops._pPr.tabs]
    except TypeError:
        return False

    if paragraph.runs[0].bold == True and any(x in tabstop_positions for x in [1962785, 2578735, 3192780]):
        return True
    return False


def has_middle_cell_multiple_codes(paragraphs) -> bool:

    if count_matching(code_condition, paragraphs) > 1:
        return True
    return False


def write_dataelement_to_dataframe(
    dataframe,
    row_index,
    edifact_struktur_cell,
    edifact_struktur_cell_left_indent_position: int,
    middle_cell,
    middle_cell_left_indent_position: int,
    tabstop_positions: List,
    bedingung_cell,
):

    # EDIFACT STRUKTUR COLUMN
    parse_paragraph_in_edifact_struktur_column_to_dataframe(
        paragraph=edifact_struktur_cell.paragraphs[0],
        dataframe=dataframe,
        row_index=row_index,
        edifact_struktur_cell_left_indent_position=edifact_struktur_cell_left_indent_position,
    )

    # BEDINGUNG COLUMN
    # Bedingung have to be parsed before MIDDLE COLUMN
    # The cell with multiple codes counts up the row_index
    # This will cause then an IndexError for Bedingung
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=dataframe, row_index=row_index)

    # MIDDLE COLUMN
    if not has_middle_cell_multiple_codes(paragraphs=middle_cell.paragraphs):
        for paragraph in middle_cell.paragraphs:
            parse_paragraph_in_middle_column_to_dataframe(
                paragraph=paragraph,
                dataframe=dataframe,
                row_index=row_index,
                left_indent_position=middle_cell_left_indent_position,
                tabstop_positions=tabstop_positions,
            )
        row_index = row_index + 1

    else:
        # The middle cell contains now multiple Codes

        # here we have to look into the next row to see, if we have to add a new datarow or
        # if we have to collect more information in the next row which we have to add to the current row

        create_new_dataframe_row_indicator_list: List = [
            paragraph.runs[0].bold == True for paragraph in middle_cell.paragraphs
        ]

        for paragraph, i in zip(middle_cell.paragraphs, range(len(create_new_dataframe_row_indicator_list))):

            # For reasons of good readability the EDIFACT Struktur information gets written again

            # EDIFACT STRUKTUR COLUMN
            parse_paragraph_in_edifact_struktur_column_to_dataframe(
                paragraph=edifact_struktur_cell.paragraphs[0],
                dataframe=dataframe,
                row_index=row_index,
                edifact_struktur_cell_left_indent_position=edifact_struktur_cell_left_indent_position,
            )

            if paragraph.runs[0].bold:
                parse_paragraph_in_middle_column_to_dataframe(
                    paragraph=paragraph,
                    dataframe=dataframe,
                    row_index=row_index,
                    left_indent_position=middle_cell_left_indent_position,
                    tabstop_positions=tabstop_positions,
                )

            elif paragraph.paragraph_format.left_indent == tabstop_positions[0]:
                # multi line Beschreibung
                dataframe.at[row_index, "Beschreibung"] += " " + paragraph.text

            if len(create_new_dataframe_row_indicator_list) > i + 1:
                if create_new_dataframe_row_indicator_list[i + 1]:
                    row_index = row_index + 1
                    dataframe.loc[row_index] = (len(dataframe.columns)) * [""]
            else:
                row_index = row_index + 1

    return row_index


def write_new_row_in_dataframe(
    row_type,
    table,
    row,
    index_for_middle_column,
    dataframe,
    dataframe_row_index,
    edifact_struktur_cell_left_indent_position,
    middle_cell_left_indent_position,
    tabstop_positions,
):
    if row_type is RowType.HEADER:
        pass

    elif row_type is RowType.SEGMENTNAME:
        write_segment_name_to_dataframe(
            dataframe=dataframe,
            row_index=dataframe_row_index,
            edifact_struktur_cell=table.row_cells(row)[0],
            edifact_struktur_cell_left_indent_position=edifact_struktur_cell_left_indent_position,
            middle_cell=table.row_cells(row)[index_for_middle_column],
            middle_cell_left_indent_position=middle_cell_left_indent_position,
            tabstop_positions=tabstop_positions,
            bedingung_cell=table.row_cells(row)[-1],
        )
        dataframe_row_index = dataframe_row_index + 1

    elif row_type is RowType.SEGMENTGRUPPE:
        write_segmentgruppe_to_dataframe(
            dataframe=dataframe,
            row_index=dataframe_row_index,
            edifact_struktur_cell=table.row_cells(row)[0],
            edifact_struktur_cell_left_indent_position=edifact_struktur_cell_left_indent_position,
            middle_cell=table.row_cells(row)[index_for_middle_column],
            middle_cell_left_indent_position=middle_cell_left_indent_position,
            tabstop_positions=tabstop_positions,
            bedingung_cell=table.row_cells(row)[-1],
        )
        dataframe_row_index = dataframe_row_index + 1

    elif row_type is RowType.SEGMENT:
        write_segment_to_dataframe(
            dataframe=dataframe,
            row_index=dataframe_row_index,
            edifact_struktur_cell=table.row_cells(row)[0],
            edifact_struktur_cell_left_indent_position=edifact_struktur_cell_left_indent_position,
            middle_cell=table.row_cells(row)[index_for_middle_column],
            middle_cell_left_indent_position=middle_cell_left_indent_position,
            tabstop_positions=tabstop_positions,
            bedingung_cell=table.row_cells(row)[-1],
        )
        dataframe_row_index = dataframe_row_index + 1

    elif row_type is RowType.DATENELEMENT:
        dataframe_row_index = write_dataelement_to_dataframe(
            dataframe=dataframe,
            row_index=dataframe_row_index,
            edifact_struktur_cell=table.row_cells(row)[0],
            edifact_struktur_cell_left_indent_position=edifact_struktur_cell_left_indent_position,
            middle_cell=table.row_cells(row)[index_for_middle_column],
            middle_cell_left_indent_position=middle_cell_left_indent_position,
            tabstop_positions=tabstop_positions,
            bedingung_cell=table.row_cells(row)[-1],
        )

    elif row_type is RowType.EMPTY:
        pass

    return dataframe_row_index
