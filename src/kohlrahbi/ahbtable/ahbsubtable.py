"""
This module contains the AhbSubTable class.
"""

from typing import Generator, Union

import numpy as np
import pandas as pd
from docx.table import Table as DocxTable
from docx.table import _Cell, _Row
from docx.text.paragraph import Paragraph
from numpy.typing import NDArray
from pydantic import BaseModel, ConfigDict

from kohlrahbi.ahbtable.ahbtablerow import AhbTableRow
from kohlrahbi.docxtablecells.bodycell import INDEX_OF_CODES_AND_QUALIFIER_COLUMN, KNOW_SUFFIXES
from kohlrahbi.enums import RowType
from kohlrahbi.row_type_checker import get_row_type
from kohlrahbi.seed import Seed


class AhbSubTable(BaseModel):
    """
    The AHB table for one Pruefidentifikator is separated into small sub tables.
    This class contains the information from such a sub table.
    """

    table_meta_data: Seed
    table: pd.DataFrame

    model_config = ConfigDict(arbitrary_types_allowed=True)

    @staticmethod
    def _parse_docx_table(
        table_meta_data: Seed, ahb_table_dataframe: pd.DataFrame, docx_table: DocxTable
    ) -> pd.DataFrame:
        """Parse the docx table and add the information to the dataframe."""
        for row in docx_table.rows:
            sanitized_cells = list(AhbSubTable.iter_visible_cells(row=row))

            current_edifact_struktur_cell = sanitized_cells[0]

            # check for row type
            current_row_type = get_row_type(
                edifact_struktur_cell=current_edifact_struktur_cell,
                left_indent_position=table_meta_data.edifact_struktur_left_indent_position,
            )

            edifact_struktur_cell = sanitized_cells[0]
            middle_cell = sanitized_cells[1]
            bedingung_cell = sanitized_cells[-1]

            # this case covers the "normal" docx table row
            if not (current_row_type is RowType.EMPTY and table_meta_data.last_two_row_types[0] is RowType.HEADER):
                ahb_table_row = AhbTableRow(
                    seed=table_meta_data,
                    edifact_struktur_cell=edifact_struktur_cell,
                    middle_cell=middle_cell,
                    bedingung_cell=bedingung_cell,
                )

                ahb_table_row_dataframe = ahb_table_row.parse(row_type=current_row_type)

                if ahb_table_row_dataframe is not None:
                    ahb_table_dataframe = pd.concat([ahb_table_dataframe, ahb_table_row_dataframe], ignore_index=True)
            else:
                # this case covers the page break situation

                # check for conditions_text
                contains_condition_texts = any(paragraph.text != "" for paragraph in bedingung_cell.paragraphs)
                # conditions are always at the top of a dataelement
                # add condition texts
                if contains_condition_texts:
                    AhbSubTable.combine_condition_text(ahb_table_dataframe, bedingung_cell)

                # add new row regularly
                ahb_table_row = AhbTableRow(
                    seed=table_meta_data,
                    edifact_struktur_cell=edifact_struktur_cell,
                    middle_cell=middle_cell,
                    bedingung_cell=bedingung_cell,
                )
                ahb_table_row_dataframe = ahb_table_row.parse(row_type=current_row_type)

                # look at first line to determine if it is broken
                first_paragraph = middle_cell.paragraphs[0]

                if ahb_table_row_dataframe is not None:
                    if AhbSubTable.is_broken_line(
                        table=ahb_table_dataframe,
                        table_meta_data=table_meta_data,
                        paragraph=first_paragraph,
                    ):
                        AhbSubTable.add_broken_line(ahb_table_dataframe, ahb_table_row_dataframe)
                        # we have a broken line
                        ahb_table_dataframe = pd.concat(
                            [ahb_table_dataframe, ahb_table_row_dataframe.iloc[1:]],
                            ignore_index=True,
                        )
                    else:
                        ahb_table_dataframe = pd.concat(
                            [ahb_table_dataframe, ahb_table_row_dataframe],
                            ignore_index=True,
                        )

            # An AhbSubTable can span over two pages.
            # But after every page break, even if we're still in the same subtable,
            # there'll be the header at the top of every page.
            # This distracts our detection logic but we workaround it by remembering
            # the last two row types.
            table_meta_data.last_two_row_types[1] = table_meta_data.last_two_row_types[0]
            table_meta_data.last_two_row_types[0] = current_row_type
        return ahb_table_dataframe

    @staticmethod
    def initialize_dataframe_with_columns(columns: list[str]) -> pd.DataFrame:
        """
        Initialize a new dataframe with the given columns
        """
        return pd.DataFrame(
            columns=columns,
            dtype="str",
        )

    @classmethod
    def from_table_with_header(cls, docx_table: DocxTable) -> "AhbSubTable":
        """
        Create a new AhbSubTable instance from a docx table WITH header
        """

        ahb_table_meta_data = Seed.from_table(docx_table=docx_table)

        ahb_table_dataframe = AhbSubTable.initialize_dataframe_with_columns(columns=ahb_table_meta_data.column_headers)

        ahb_table_dataframe = cls._parse_docx_table(
            table_meta_data=ahb_table_meta_data,
            ahb_table_dataframe=ahb_table_dataframe,
            docx_table=docx_table,
        )

        return cls(table_meta_data=ahb_table_meta_data, table=ahb_table_dataframe)

    @classmethod
    def from_headless_table(cls, tmd: Seed, docx_table: DocxTable) -> "AhbSubTable":
        """
        Create a new AhbSubTable instance from a docx table WITHOUT header
        """

        ahb_table_dataframe = AhbSubTable.initialize_dataframe_with_columns(columns=tmd.column_headers)

        ahb_table_dataframe = cls._parse_docx_table(
            table_meta_data=tmd,
            ahb_table_dataframe=ahb_table_dataframe,
            docx_table=docx_table,
        )

        return cls(table_meta_data=tmd, table=ahb_table_dataframe)

    @staticmethod
    def iter_visible_cells(row: _Row) -> Generator[_Cell, None, None]:
        """
        This function makes sure that you will iterate over the cells you see in the word document.
        For more information go to https://github.com/python-openxml/python-docx/issues/970#issuecomment-877386927
        """
        table_row = row._tr  # pylint:disable=protected-access
        for table_column in table_row.tc_lst:
            yield _Cell(table_column, row.table)

    @staticmethod
    def add_text_to_last_row(ahb_table_dataframe: pd.DataFrame, row_index: int, column_index: int, text: str) -> None:
        """Add a text to the last row of the dataframe."""
        starts_with_known_suffix = any(text.startswith(suffix + " ") for suffix in KNOW_SUFFIXES)
        if len(text) > 0:
            if len(ahb_table_dataframe.iat[row_index, column_index]) > 0 and not starts_with_known_suffix:
                text = " " + text
            ahb_table_dataframe.iat[row_index, column_index] += text

    @staticmethod
    def add_broken_line(ahb_table_dataframe: pd.DataFrame, broken_line: pd.DataFrame) -> None:
        """Add a broken line to the dataframe."""
        for col_index in range(INDEX_OF_CODES_AND_QUALIFIER_COLUMN, len(ahb_table_dataframe.columns)):
            AhbSubTable.add_text_to_last_row(
                ahb_table_dataframe, ahb_table_dataframe.index.max(), col_index, str(broken_line.iat[0, col_index])
            )

    @staticmethod
    def combine_condition_text(ahb_table_dataframe: pd.DataFrame, bedingung_cell: _Cell) -> None:
        """Add the condition text to the dataframe."""
        conditions_text = " " + " ".join(
            paragraph.text for paragraph in bedingung_cell.paragraphs if paragraph.text != ""
        )
        last_valid_row = ahb_table_dataframe["Bedingung"].last_valid_index()
        conditions_text = ahb_table_dataframe.at[last_valid_row, "Bedingung"] + conditions_text
        # remove existing text
        ahb_table_dataframe.at[last_valid_row, "Bedingung"] = ""
        # remove remaining text to avoid misplacements
        for paragraph in bedingung_cell.paragraphs:
            paragraph.text = ""
        bedingung_cell.paragraphs[-1].text = conditions_text

    @staticmethod
    def is_broken_line(
        table: pd.DataFrame,
        table_meta_data: Seed,
        paragraph: Paragraph,
    ) -> bool:
        """
        Check for broken lines in the middle cell.
        """
        tabsplit_text = paragraph.text.split("\t")

        loc: Union[int, slice, NDArray[np.bool_]] = table.columns.get_loc("Beschreibung")

        # Ensure loc is an int
        if isinstance(loc, int):
            beschreibung_index: int = loc
        else:
            raise ValueError("The location of the column 'Beschreibung' is not an integer.")

        is_empty_middle_line = all(text == "" for text in tabsplit_text)
        is_broken_code_qualifier = (
            paragraph.paragraph_format.left_indent is not None
            and paragraph.paragraph_format.left_indent != table_meta_data.middle_cell_left_indent_position
            and table.iat[-1, beschreibung_index] != ""
            and table.iloc[-1, beschreibung_index + 1 :].ne("").any()
        )
        if is_broken_code_qualifier and len(tabsplit_text) == 1:
            # only broken code / qualifier
            assert (
                table.iat[-1, beschreibung_index] != "" and table.iloc[-1, beschreibung_index + 1 :].ne("").any()
            ), "no condition expected in broken line"
        there_are_conditions = (
            len(tabsplit_text) > 1
            and paragraph.paragraph_format.left_indent != table_meta_data.middle_cell_left_indent_position
        )

        return is_empty_middle_line or there_are_conditions or is_broken_code_qualifier
