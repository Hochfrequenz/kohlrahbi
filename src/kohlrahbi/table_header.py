"""
This module contains the TableHeader class.
"""

from enum import StrEnum
from typing import Dict, List, Mapping, cast

from docx.table import _Cell
from docx.text.paragraph import Paragraph
from more_itertools import first, last
from pydantic import BaseModel, Field


class HeaderSection(StrEnum):
    """
    Enum for the different sections of the table header
    """

    BESCHREIBUNG = "beschreibung"
    KOMMUNIKATION_VON = "kommunikation von"
    PRUEFIDENTIFIKATOR = "pruefidentifikator"


def get_tabstop_positions(paragraph: Paragraph) -> List[int]:
    """Find all tabstop positions in a given paragraph.

    Mainly the tabstop positions of cells from the middle column are determined

    Args:
        paragraph (Paragraph):

    Returns:
        List[int]: All tabstop positions in the given paragraph
    """
    tabstop_positions: List[int] = []
    for tabstop in paragraph.paragraph_format.tab_stops:
        tabstop_positions.append(tabstop.position)
    return tabstop_positions


def create_mapping_of_tabstop_positions(
    initial_tabstop_positions: List[int],
    current_tabstop_positions: List[int],
) -> Mapping[int, int]:
    """
    Create a mapping of the tabstop positions of the Prüfidentifikatoren columns.

    For the current tabstop positions, the intial ones with the least difference are
    found to account for changes in tabstop positions between paragraphs.

    Returns:
        Dict[int, int]: All initial tabstop positions mapped to the current tabstop positions
    """
    # create a mapping of the tabstop positions
    mapping: Dict[int, int] = {}
    # Sort the lists in ascending order
    initial_tabstop_positions.sort()
    current_tabstop_positions.sort()

    # Iterate over the sorted lists and find the entries with the least difference
    for i, current_pos in enumerate(current_tabstop_positions):
        min_diff = float("inf")
        min_j = None
        for j, initial_pos in enumerate(initial_tabstop_positions):
            diff = abs(current_pos - initial_pos)
            if diff < min_diff:
                min_diff = diff
                min_j = j
        if min_j is None:
            raise ValueError("min_j should not be None")

        mapping[current_tabstop_positions[i]] = initial_tabstop_positions[min_j]

    return mapping


# pylint: disable=too-few-public-methods


class PruefiMetaData(BaseModel):
    """
    This class contains the information about the Prüfidentifikatoren
    """

    pruefidentifikator: str
    name: str
    communication_direction: str


class TableHeader(BaseModel):
    """
    Class for the table header.
    It contains the information about the Prüfidentifikatoren.
    """

    pruefi_meta_data: List[PruefiMetaData] = Field(default_factory=list)

    @classmethod
    def from_header_cell(cls, row_cell: _Cell) -> "TableHeader":
        """
        Create a TableHeader instance from a list of strings.
        """

        if not row_cell.paragraphs[-1].text.startswith("Prüfidentifikator"):
            raise ValueError("The last paragraph should start with 'Prüfidentifikator'")

        collector = cls.initialize_collector(paragraph=last(row_cell.paragraphs))

        section_type: HeaderSection

        # the main goal of this loop is to find the tabstop positions of each Prüfidentifikator.
        for paragraph in row_cell.paragraphs:
            splitted_text = paragraph.text.split("\t")
            text_prefix = first(splitted_text)

            # the paragraph with the Prüfidentifikatoren has slightly different tabstop positions compared to
            # the table content rows. So we skip this paragraph.
            does_paragraph_contain_pruefidentifikatoren = text_prefix == "Prüfidentifikator"
            if does_paragraph_contain_pruefidentifikatoren:
                continue

            # only the paragraphs with the Beschreibung and Kommunikation von have the same tabstop positions as
            # the table content rows.
            does_paragraph_contain_beschreibung_or_kommunikation_von = text_prefix in {
                "Beschreibung",
                "Kommunikation von",
            }
            if does_paragraph_contain_beschreibung_or_kommunikation_von:
                initial_tabstop_positions = get_tabstop_positions(paragraph=paragraph)
                tabstop_mapper = dict(zip(initial_tabstop_positions, collector.keys()))
                section_type = HeaderSection[text_prefix.replace(" ", "_").upper()]
                splitted_text.pop(0)
                for pruefi, text in zip(collector.keys(), splitted_text):
                    collector[pruefi][section_type.value] = text + " "
            else:
                if "" in splitted_text and len(splitted_text) > len(collector.keys()):
                    splitted_text.remove("")

                for tabstop_position, text in zip(initial_tabstop_positions, splitted_text):
                    pruefi = tabstop_mapper[tabstop_position]
                    collector[pruefi][section_type.value] += text + " "

        pruefi_meta_data = [
            PruefiMetaData(
                pruefidentifikator=pruefi,
                communication_direction=cls.ensure_single_space_between_words(
                    #  The cast function is a no-op at runtime and doesn't perform any actual type conversion;
                    #  it's only used for type checking purposes.
                    cast(str, meta_data[HeaderSection.KOMMUNIKATION_VON.value]),
                ),
                name=cls.ensure_single_space_between_words(cast(str, meta_data[HeaderSection.BESCHREIBUNG.value])),
            )
            for pruefi, meta_data in collector.items()
            if isinstance(meta_data[HeaderSection.KOMMUNIKATION_VON.value], str)
        ]

        return cls(pruefi_meta_data=pruefi_meta_data)

    @staticmethod
    def initialize_collector(paragraph: Paragraph) -> Dict[str, Dict[str, str | int]]:
        """Initialize the collector"""
        current_tabstop_positions = get_tabstop_positions(paragraph=paragraph)
        splitted_text = paragraph.text.split("\t")
        splitted_text.remove("Prüfidentifikator")

        collector: Dict[str, Dict[str, str | int]] = {
            pruefidentifikator: {
                HeaderSection.BESCHREIBUNG.value: "",
                HeaderSection.KOMMUNIKATION_VON.value: "",
                "tabstop_position": tab_stop,
            }
            for pruefidentifikator, tab_stop in zip(splitted_text, current_tabstop_positions)
        }

        if not collector:
            raise ValueError("collector should not be empty")

        return collector

    @staticmethod
    def ensure_single_space_between_words(text: str) -> str:
        """Inserts a single space between words"""
        return " ".join(text.split())

    def get_pruefidentifikatoren(self) -> List[str]:
        """
        Get all Prüfidentifikatoren from the table header.
        The order of the Prüfidentifikatoren is the same as in the docx table headers.
        So there should be no duplicates.
        """
        # pylint:disable=not-an-iterable
        return [pruefi.pruefidentifikator for pruefi in self.pruefi_meta_data]
