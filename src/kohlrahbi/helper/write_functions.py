"""
Collection of functions to write the extracted infos from the AHB tables into a DataFrame.
"""


from docx.table import Table, _Cell  # type:ignore[import]

from kohlrahbi.helper.row_type_checker import RowType
from kohlrahbi.helper.seed import Seed
from kohlrahbi.parser import parse_bedingung_cell, parse_edifact_struktur_cell, parse_middle_cell


# pylint: disable=too-many-arguments
def write_segment_name_to_dataframe(
    seed: Seed,
    edifact_struktur_cell: _Cell,
    middle_cell: _Cell,
    bedingung_cell: _Cell,
) -> None:
    """Writes all infos from a segment name row into a DataFrame.

    Args:
        seed (Seed):
        edifact_struktur_cell (_Cell): Cell from the edifact struktur column
        middle_cell (_Cell): Cell from the middle column
        bedingung_cell (_Cell): Cell from the Bedingung column
    """

    seed.soul.loc[seed.soul.index.max() + 1, :] = ""

    # EDIFACT STRUKTUR COLUMN
    parse_edifact_struktur_cell(
        table_cell=edifact_struktur_cell,
        dataframe=seed.soul,
        row_index=seed.current_df_row_index,
        edifact_struktur_cell_left_indent_position=seed.edifact_struktur_left_indent_position,
    )

    # MIDDLE COLUMN
    # I do not expect to get a multiline Segmentgruppe,
    # but just in case we loop through all paragraphs

    parse_middle_cell(
        table_cell=middle_cell,
        dataframe=seed.soul,
        row_index=seed.current_df_row_index,
        left_indent_position=seed.middle_cell_left_indent_position,
        indicator_tabstop_positions=seed.tabstop_positions,
    )

    # BEDINGUNG COLUMN
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=seed.soul, row_index=seed.current_df_row_index)


# pylint: disable=too-many-arguments
def write_segmentgruppe_to_dataframe(
    elixir: Seed,
    edifact_struktur_cell: _Cell,
    middle_cell: _Cell,
    bedingung_cell: _Cell,
) -> None:
    """Writes all infos from a segmentgruppe row into a DataFrame.

    Args:
        elixir (Elixir):
        edifact_struktur_cell (_Cell): Cell from the edifact struktur column
        middle_cell (_Cell): Cell from the middle column
        bedingung_cell (_Cell): Cell from the Bedingung column
    """

    elixir.soul.loc[elixir.soul.index.max() + 1, :] = ""

    # EDIFACT STRUKTUR COLUMN
    parse_edifact_struktur_cell(
        # there might be 2 paragraphs in case of multi line headings, so we're handing over all the paragraphs
        table_cell=edifact_struktur_cell,
        dataframe=elixir.soul,
        row_index=elixir.current_df_row_index,
        edifact_struktur_cell_left_indent_position=elixir.edifact_struktur_left_indent_position,
    )

    # MIDDLE COLUMN
    # I do not expect to a multiline Segementgruppe,
    # but just in case we loop through all paragraphs
    parse_middle_cell(
        table_cell=middle_cell,
        dataframe=elixir.soul,
        row_index=elixir.current_df_row_index,
        left_indent_position=elixir.middle_cell_left_indent_position,
        indicator_tabstop_positions=elixir.tabstop_positions,
    )

    # BEDINGUNG COLUMN
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=elixir.soul, row_index=elixir.current_df_row_index)


# pylint: disable=too-many-arguments
def write_segment_to_dataframe(
    elixir: Seed,
    edifact_struktur_cell: _Cell,
    middle_cell: _Cell,
    bedingung_cell: _Cell,
) -> None:
    """Writes all infos from a segment row into a DataFrame.

    Args:
        elixir (Elixir):
        edifact_struktur_cell (_Cell): Cell from the edifact struktur column
        middle_cell (_Cell): Cell from the middle column
        bedingung_cell (_Cell): Cell from the Bedingung column
    """

    elixir.soul.loc[elixir.soul.index.max() + 1, :] = ""

    # EDIFACT STRUKTUR COLUMN
    parse_edifact_struktur_cell(
        table_cell=edifact_struktur_cell,
        dataframe=elixir.soul,
        row_index=elixir.current_df_row_index,
        edifact_struktur_cell_left_indent_position=elixir.edifact_struktur_left_indent_position,
    )

    # MIDDLE COLUMN
    parse_middle_cell(
        table_cell=middle_cell,
        dataframe=elixir.soul,
        row_index=elixir.current_df_row_index,
        left_indent_position=elixir.middle_cell_left_indent_position,
        indicator_tabstop_positions=elixir.tabstop_positions,
    )

    # BEDINGUNG COLUMN
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=elixir.soul, row_index=elixir.current_df_row_index)


# pylint: disable=too-many-arguments
def write_dataelement_to_dataframe(
    elixir: Seed,
    edifact_struktur_cell: _Cell,
    middle_cell: _Cell,
    bedingung_cell: _Cell,
) -> int:
    """Writes all infos from a dataelement row into a DataFrame.

    Args:
        elixir (Elixir):
        edifact_struktur_cell (_Cell): Cell from the edifact struktur column
        middle_cell (_Cell): Cell from the middle column
        bedingung_cell (_Cell): Cell from the Bedingung column
    """

    elixir.soul.loc[elixir.soul.index.max() + 1, :] = ""

    # EDIFACT STRUKTUR COLUMN
    parse_edifact_struktur_cell(
        table_cell=edifact_struktur_cell,
        dataframe=elixir.soul,
        row_index=elixir.current_df_row_index,
        edifact_struktur_cell_left_indent_position=elixir.edifact_struktur_left_indent_position,
    )

    # BEDINGUNG COLUMN
    # Bedingung have to be parsed before MIDDLE COLUMN
    # The cell with multiple codes counts up the row_index
    # This will cause then an IndexError for Bedingung
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=elixir.soul, row_index=elixir.current_df_row_index)

    # GOAL
    parse_middle_cell(
        table_cell=middle_cell,
        dataframe=elixir.soul,
        row_index=elixir.current_df_row_index,
        left_indent_position=elixir.middle_cell_left_indent_position,
        indicator_tabstop_positions=elixir.tabstop_positions,
    )

    # MIDDLE COLUMN
    # if not has_middle_cell_multiple_codes(
    #     paragraphs=middle_cell.paragraphs, pruefi_tabstops=elixir.tabstop_positions[1:]
    # ):
    #     # for paragraph in middle_cell.paragraphs:
    #     parse_middle_cell(
    #         table_cell=middle_cell,
    #         dataframe=elixir.soul,
    #         row_index=elixir.current_df_row_index,
    #         left_indent_position=elixir.middle_cell_left_indent_position,
    #         tabstop_positions=elixir.tabstop_positions,
    #     )
    #     elixir.current_df_row_index = elixir.current_df_row_index + 1

    # else:
    #     # The middle cell contains multiple Codes

    #     # here we have to look into the next row to see, if we have to add a new datarow or
    #     # if we have to collect more information in the next row which we have to add to the current row

    #     create_new_dataframe_row_indicator_list: List = [
    #         paragraph.runs[0].bold is True for paragraph in middle_cell.paragraphs
    #     ]

    #     for paragraph, i in zip(middle_cell.paragraphs, range(len(create_new_dataframe_row_indicator_list))):

    #         # For reasons of good readability the EDIFACT Struktur information gets written again

    #         # EDIFACT STRUKTUR COLUMN

    #         if edifact_struktur_cell.paragraphs[0].text != "":
    #             parse_edifact_struktur_cell(
    #                 table_cell=edifact_struktur_cell,
    #                 dataframe=elixir.soul,
    #                 row_index=elixir.current_df_row_index,
    #                 edifact_struktur_cell_left_indent_position=elixir.edifact_struktur_left_indent_position,
    #             )
    #         else:
    #             elixir.soul.at[elixir.current_df_row_index, "Segment Gruppe"] = elixir.soul.loc[
    #                 elixir.current_df_row_index - 1, "Segment Gruppe"
    #             ]
    #             elixir.soul.at[elixir.current_df_row_index, "Segment"] = elixir.soul.loc[
    #                 elixir.current_df_row_index - 1, "Segment"
    #             ]
    #             elixir.soul.at[elixir.current_df_row_index, "Datenelement"] = elixir.soul.loc[
    #                 elixir.current_df_row_index - 1, "Datenelement"
    #             ]

    #         if paragraph.runs[0].bold:
    #             parse_middle_cell(
    #                 paragraph=paragraph,
    #                 dataframe=elixir.soul,
    #                 row_index=elixir.current_df_row_index,
    #                 left_indent_position=elixir.middle_cell_left_indent_position,
    #                 tabstop_positions=elixir.tabstop_positions,
    #             )

    #         elif paragraph.paragraph_format.left_indent == elixir.tabstop_positions[0]:
    #             # multi line Beschreibung
    #             elixir.soul.at[elixir.current_df_row_index, "Beschreibung"] += " " + paragraph.text

    #         if len(create_new_dataframe_row_indicator_list) > i + 1:
    #             if create_new_dataframe_row_indicator_list[i + 1]:
    #                 elixir.current_df_row_index = elixir.current_df_row_index + 1
    #                 elixir.soul.loc[elixir.current_df_row_index] = (len(elixir.soul.columns)) * [""]
    #         else:
    #             elixir.current_df_row_index = elixir.current_df_row_index + 1

    return elixir.current_df_row_index


def write_new_row_in_dataframe(
    elixir: Seed,
    row_type: RowType,
    table: Table,
    row: int,
    index_for_middle_column: int,
) -> None:
    """Writes the current row of the current table into the DataFrame depending on the type of the row

    Args:
        elixir (Elixir):
        row_type (RowType): Type of the current row
        table (Table): Current table
        row (int): Row of the current table
        index_for_middle_column (int): Index of the actual middle column
    """

    if row_type is RowType.HEADER:
        pass

    elif row_type is RowType.SEGMENTNAME:
        write_segment_name_to_dataframe(
            seed=elixir,
            edifact_struktur_cell=table.row_cells(row)[0],
            middle_cell=table.row_cells(row)[index_for_middle_column],
            bedingung_cell=table.row_cells(row)[-1],
        )

    elif row_type is RowType.SEGMENTGRUPPE:
        write_segmentgruppe_to_dataframe(
            elixir=elixir,
            edifact_struktur_cell=table.row_cells(row)[0],
            middle_cell=table.row_cells(row)[index_for_middle_column],
            bedingung_cell=table.row_cells(row)[-1],
        )

    elif row_type is RowType.SEGMENT:
        write_segment_to_dataframe(
            elixir=elixir,
            edifact_struktur_cell=table.row_cells(row)[0],
            middle_cell=table.row_cells(row)[index_for_middle_column],
            bedingung_cell=table.row_cells(row)[-1],
        )

    elif row_type is RowType.DATENELEMENT:
        write_dataelement_to_dataframe(
            elixir=elixir,
            edifact_struktur_cell=table.row_cells(row)[0],
            middle_cell=table.row_cells(row)[index_for_middle_column],
            bedingung_cell=table.row_cells(row)[-1],
        )

    elif row_type is RowType.EMPTY:
        pass
