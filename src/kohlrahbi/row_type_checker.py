"""
This module contains all functions to define the type of a row of the tables in an AHB.
"""
from docx.oxml.ns import qn  # type:ignore[import]
from docx.oxml.parser import OxmlElement  # type:ignore[import]
from docx.shared import RGBColor  # type:ignore[import]
from docx.table import _Cell  # type:ignore[import]

from kohlrahbi.enums import RowType


def set_table_header_bg_color(cell, hex_color: str):
    """
    set background shading for Header Rows
    """
    table_cell = cell._tc  # pylint:disable=protected-access
    table_cell_properties = table_cell.get_or_add_tcPr()
    cl_shading = OxmlElement("w:shd")
    cl_shading.set(qn("w:fill"), hex_color)  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    table_cell_properties.append(cl_shading)
    return cell


def is_row_header(edifact_struktur_cell: _Cell) -> bool:
    """Checks if the current row is a header.

    Args:
        edifact_struktur_cell (_Cell): Indicator cell

    Returns:
        bool:
    """
    if edifact_struktur_cell.text == "EDIFACT Struktur":
        return True

    return False


def is_row_segmentname(edifact_struktur_cell: _Cell) -> bool:
    """Checks if the current row contains just a segment name.
       Example: "Nachrichten-Kopfsegment"

    Args:
        edifact_struktur_cell (_Cell): Indicator cell

    Returns:
        bool:
    """
    try:
        return edifact_struktur_cell.paragraphs[0].runs[0].font.color.rgb == RGBColor(128, 128, 128)  # grey
    except IndexError:
        return False


def is_row_segmentgruppe(edifact_struktur_cell: _Cell, left_indent_position: int) -> bool:
    """Checks if the current row is a segmentgruppe.
       Example: "SG2"

    Args:
        edifact_struktur_cell (_Cell): Indicator cell
        left_indent_position (int): Position of the left indent

    Returns:
        bool:
    """
    return (
        edifact_struktur_cell.paragraphs[0].paragraph_format.left_indent != left_indent_position
        and "\t" not in edifact_struktur_cell.text
        and not edifact_struktur_cell.text == ""
    )


def is_row_segment(edifact_struktur_cell: _Cell, left_indent_position: int) -> bool:
    """Checks if the current row is a segment.
       Example: "UNH", "SG2\tNAD"

    Args:
        edifact_struktur_cell (_Cell): Indicator cell
        left_indent_position (int): Position of the left indent

    Returns:
        bool:
    """
    # |   UNH    |
    if (
        edifact_struktur_cell.paragraphs[0].paragraph_format.left_indent == left_indent_position
        and "\t" not in edifact_struktur_cell.text
        and not edifact_struktur_cell.text == ""
    ):
        return True

    # | SG2\tNAD |
    if (
        not edifact_struktur_cell.paragraphs[0].paragraph_format.left_indent == left_indent_position
        and edifact_struktur_cell.text.count("\t") == 1
    ):
        return True

    return False


def is_row_datenelement(edifact_struktur_cell: _Cell, left_indent_position: int) -> bool:
    """Checks if the current row is a datenelement.
       Example: "UNH\t00062", "SG2\tNAD\t3035"

    Args:
        edifact_struktur_cell (_Cell): Indicator cell
        left_indent_position (int): Position of the left indent

    Returns:
        bool:
    """
    # |   UNH\t0062 |
    if (
        edifact_struktur_cell.paragraphs[0].paragraph_format.left_indent == left_indent_position
        and "\t" in edifact_struktur_cell.text
    ):
        return True

    # | SG2\tNAD\t3035 |
    if (
        not edifact_struktur_cell.paragraphs[0].paragraph_format.left_indent == left_indent_position
        and edifact_struktur_cell.text.count("\t") == 2
    ):
        return True

    return False


def is_row_empty(edifact_struktur_cell: _Cell) -> bool:
    """Checks if the current row is empty.
       Example: ""
    Args:
        edifact_struktur_cell (_Cell): Indicator cell

    Returns:
        bool:
    """
    return edifact_struktur_cell.text == ""


def get_row_type(edifact_struktur_cell: _Cell, left_indent_position: int) -> RowType:
    """Defines the type of the current row.

    Args:
        edifact_struktur_cell (_Cell): Indicator cell
        left_indent_position (int): Position of the left indent

    Raises:
        NotImplemented: Gets raised if the RowType got not to be defined

    Returns:
        RowType: Type of the current row
    """
    if is_row_header(edifact_struktur_cell=edifact_struktur_cell):
        return RowType.HEADER

    if is_row_segmentname(edifact_struktur_cell=edifact_struktur_cell):
        return RowType.SEGMENTNAME

    if is_row_segmentgruppe(edifact_struktur_cell=edifact_struktur_cell, left_indent_position=left_indent_position):
        return RowType.SEGMENTGRUPPE

    if is_row_segment(edifact_struktur_cell=edifact_struktur_cell, left_indent_position=left_indent_position):
        return RowType.SEGMENT

    if is_row_datenelement(edifact_struktur_cell=edifact_struktur_cell, left_indent_position=left_indent_position):
        return RowType.DATENELEMENT

    if is_row_empty(edifact_struktur_cell=edifact_struktur_cell):
        return RowType.EMPTY

    raise NotImplementedError(f"Could not define row type of cell with text: {edifact_struktur_cell.text}")
