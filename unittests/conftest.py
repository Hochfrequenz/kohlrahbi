from collections.abc import Callable

import docx
import docx.table
import pytest

from unittests.cellparagraph import CellParagraph


@pytest.fixture(autouse=True)
def _disable_rich_color(monkeypatch: pytest.MonkeyPatch) -> None:
    """
    Force rich/typer CLI output to render without ANSI escape codes.
    Some environments (e.g. CI runners) set FORCE_COLOR, which makes rich treat output as a
    color-capable terminal even though typer.testing.CliRunner captures a non-tty stream. Rich
    then emits bold/dim styling around option names, splitting e.g. "--url" into separately
    escaped spans ("-" + "-url") and breaking plain substring assertions against CliRunner
    output. NO_COLOR alone does not prevent this, since bold/dim are style, not color - the
    terminal detection itself (driven by FORCE_COLOR) has to be neutralized.
    """
    monkeypatch.delenv("FORCE_COLOR", raising=False)
    monkeypatch.setenv("NO_COLOR", "1")


@pytest.fixture
def get_ahb_table_with_multiple_paragraphs() -> Callable[[list[CellParagraph]], docx.table.Table]:
    def _setup_ahb_table(body_cell_paragraphs: list[CellParagraph]) -> docx.table.Table:
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)

        body_cell = table.rows[0].cells[0]

        # the cell comes with an empty paragraph which I could not delete.
        # So we insert the BodyCellParagraph attributes into the empty paragraph
        first_body_cell_paragprah: CellParagraph = body_cell_paragraphs[0]

        body_cell.paragraphs[0].text = first_body_cell_paragprah.text

        if first_body_cell_paragprah.tabstop_positions is not None:
            for tabstop_position in first_body_cell_paragprah.tabstop_positions:
                body_cell.paragraphs[0].paragraph_format.tab_stops.add_tab_stop(tabstop_position)

        body_cell.paragraphs[0].paragraph_format.left_indent = first_body_cell_paragprah.left_indent_length

        for paragraph_index, bcp in zip(range(1, len(body_cell_paragraphs[1:]) + 1), body_cell_paragraphs[1:]):
            # add paragraph with text
            body_cell.add_paragraph(text=bcp.text)

            # set tabstop positions
            if bcp.tabstop_positions is not None:
                for tabstop_position in bcp.tabstop_positions:
                    body_cell.paragraphs[paragraph_index].paragraph_format.tab_stops.add_tab_stop(tabstop_position)

            # set left indent lenght
            body_cell.paragraphs[paragraph_index].paragraph_format.left_indent = bcp.left_indent_length

        result: docx.table.Table = table
        return result

    return _setup_ahb_table
