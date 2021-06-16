from dataclasses import dataclass
from typing import List

import docx
import pandas as pd
import pytest

from ahb_extractor.helper.write_functions import (
    parse_bedingung_cell,
    parse_paragraph_in_edifact_struktur_column_to_dataframe,
    parse_paragraph_in_middle_column_to_dataframe,
    write_dataelement_to_dataframe,
    write_segment_name_to_dataframe,
    write_segment_to_dataframe,
    write_segmentgruppe_to_dataframe,
)


class TestParseFunctions:

    # create table test cell
    # it contains per default an empty paragraph
    test_document = docx.Document()
    test_table = test_document.add_table(rows=1, cols=1)
    test_cell = test_table.add_row().cells[0]

    # this left indent and tabstop positions are equal to
    # the left indent and tabstop positions of the indicator paragraph
    middle_cell_left_indent_position_of_indicator_paragraph = 36830
    middle_cell_tabstop_positions_of_indicator_paragraph = [436245, 1962785, 2578735, 3192780]

    edifact_struktur_cell_left_indent_position_of_segmentgroup_cells = 36830
    edifact_struktur_cell_left_indent_position_of_indicator_paragraph = 364490

    @pytest.mark.parametrize(
        "text_content, left_indent_position, cell_tabstop_positions, expected_df_row",
        [
            pytest.param(
                "",
                None,
                None,
                {
                    "Segment Gruppe": "",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="Segmentname Nachrichten-Kopfsegment",
            ),
            pytest.param(
                "\tMuss\tMuss\tMuss",
                None,
                middle_cell_tabstop_positions_of_indicator_paragraph,
                {
                    "Segment Gruppe": "",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "Muss",
                    "88888": "Muss",
                    "99999": "Muss",
                    "Bedingung": "",
                },
                id="First UNH Segment",
            ),
            pytest.param(
                "Nachrichten-Referenznummer\tX\tX\tX",
                middle_cell_left_indent_position_of_indicator_paragraph,
                middle_cell_tabstop_positions_of_indicator_paragraph[1:],
                {
                    "Segment Gruppe": "",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "Nachrichten-Referenznummer",
                    "Beschreibung": "",
                    "77777": "X",
                    "88888": "X",
                    "99999": "X",
                    "Bedingung": "",
                },
                id="Qualifier",
            ),
            pytest.param(
                "UTILM\tNetzanschluss-\tX\tX\tX",
                middle_cell_left_indent_position_of_indicator_paragraph,
                middle_cell_tabstop_positions_of_indicator_paragraph,
                {
                    "Segment Gruppe": "",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "UTILM",
                    "Beschreibung": "Netzanschluss-",
                    "77777": "X",
                    "88888": "X",
                    "99999": "X",
                    "Bedingung": "",
                },
                id="First Dataelement",
            ),
            pytest.param(
                "D\tStammdaten",
                middle_cell_left_indent_position_of_indicator_paragraph,
                middle_cell_tabstop_positions_of_indicator_paragraph[0:1],
                {
                    "Segment Gruppe": "",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "D",
                    "Beschreibung": "Stammdaten",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="Dataelement without conditions",
            ),
            pytest.param(
                "zugrundeliegenden",
                middle_cell_tabstop_positions_of_indicator_paragraph[0],
                None,
                {
                    "Segment Gruppe": "",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "zugrundeliegenden",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="only Beschreibung",
            ),
            pytest.param(
                "\tMuss [16] U\n",
                None,
                middle_cell_tabstop_positions_of_indicator_paragraph[-1:],
                {
                    "Segment Gruppe": "",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "Muss [16] U\n",
                    "Bedingung": "",
                },
                id="Just one entry for one Prüfidentifikator",
            ),
        ],
    )
    def test_parse_paragraph_in_middle_column(
        self, text_content, left_indent_position, cell_tabstop_positions, expected_df_row
    ):

        # insert text
        self.test_cell.text = text_content
        test_paragraph = self.test_cell.paragraphs[0]

        # set left indent positon
        test_paragraph.paragraph_format.left_indent = left_indent_position

        tab_stops = test_paragraph.paragraph_format.tab_stops
        # Length: https://python-docx.readthedocs.io/en/latest/api/shared.html#docx.shared.Length

        if cell_tabstop_positions is not None:
            for tabstop_position in cell_tabstop_positions:
                tab_stops.add_tab_stop(tabstop_position)

        # Initial two dataframes ...
        df = pd.DataFrame(columns=expected_df_row.keys(), dtype="str")
        expected_df = pd.DataFrame(columns=expected_df_row.keys(), dtype="str")
        row_index = 0
        # ... with a row full of emtpy strings
        initial_dataframe_row = (len(df.columns)) * [""]
        df.loc[row_index] = initial_dataframe_row
        expected_df.loc[row_index] = initial_dataframe_row

        parse_paragraph_in_middle_column_to_dataframe(
            paragraph=test_paragraph,
            dataframe=df,
            row_index=row_index,
            left_indent_position=self.middle_cell_left_indent_position_of_indicator_paragraph,
            tabstop_positions=self.middle_cell_tabstop_positions_of_indicator_paragraph,
        )

        expected_df.loc[row_index] = expected_df_row

        assert expected_df.equals(df)

    # def test_not_implemented_middle_cell_paragraph(self):
    #     # insert text
    #     self.test_cell.text = ""
    #     test_paragraph = self.test_cell.paragraphs[0]

    #     # set left indent positon
    #     test_paragraph.paragraph_format.left_indent = None

    #     df = pd.DataFrame(dtype="str")
    #     row_index = 0

    #     with pytest.raises(NotImplementedError) as excinfo:
    #         parse_paragraph_in_middle_column_to_dataframe(
    #             paragraph=test_paragraph,
    #             dataframe=df,
    #             row_index=row_index,
    #             left_indent_position=self.left_indent_position_of_indicator_paragraph,
    #             tabstop_positions=self.tabstop_positions_of_indicator_paragraph,
    #         )

    #     assert "Could not parse paragraphe in middle cell with " in str(excinfo.value)

    @pytest.mark.parametrize(
        "text_content, left_indent_position, expected_df_row",
        [
            pytest.param(
                "Nachrichten-Kopfsegment",
                edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                {
                    "Segment Gruppe": "Nachrichten-Kopfsegment",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="Segmentname",
            ),
            pytest.param(
                "UNH",
                edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
                {
                    "Segment Gruppe": "",
                    "Segment": "UNH",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="Segment",
            ),
            pytest.param(
                "UNH\t0062",
                edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
                {
                    "Segment Gruppe": "",
                    "Segment": "UNH",
                    "Datenelement": "0062",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="Segment with Dataelement",
            ),
            pytest.param(
                "SG2",
                edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                {
                    "Segment Gruppe": "SG2",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="Segmentgruppe",
            ),
            pytest.param(
                "SG2\tNAD",
                edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                {
                    "Segment Gruppe": "SG2",
                    "Segment": "NAD",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="Segmentgruppe with Segment",
            ),
            pytest.param(
                "SG2\tNAD\t3035",
                edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                {
                    "Segment Gruppe": "SG2",
                    "Segment": "NAD",
                    "Datenelement": "3035",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="Segmentgruppe with Segment and Datenelement",
            ),
        ],
    )
    def test_parse_paragraph_in_edifact_struktur_column_to_dataframe(
        self, text_content, left_indent_position, expected_df_row
    ):

        # insert text
        self.test_cell.text = text_content
        test_paragraph = self.test_cell.paragraphs[0]

        # set left indent positon
        test_paragraph.paragraph_format.left_indent = left_indent_position

        # Initial two dataframes ...
        df = pd.DataFrame(columns=expected_df_row.keys(), dtype="str")
        expected_df = pd.DataFrame(columns=expected_df_row.keys(), dtype="str")
        row_index = 0
        # ... with a row full of emtpy strings
        initial_dataframe_row = (len(df.columns)) * [""]
        df.loc[row_index] = initial_dataframe_row
        expected_df.loc[row_index] = initial_dataframe_row

        parse_paragraph_in_edifact_struktur_column_to_dataframe(
            paragraph=test_paragraph,
            dataframe=df,
            row_index=row_index,
            edifact_struktur_cell_left_indent_position=self.edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
        )

        expected_df.loc[row_index] = expected_df_row

        assert expected_df.equals(df)

    @pytest.mark.parametrize(
        "text_content, expected_df_row",
        [
            pytest.param(
                "",
                {
                    "Segment Gruppe": "",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": "",
                },
                id="Empty Bedingung",
            ),
            pytest.param(
                """[12] Wenn SG4
DTM+471 (Ende zum
nächstmöglichem
Termin) nicht vorhanden

[13] Wenn SG4
STS+E01++Z01 (Status
der Antwort: Zustimmung
mit Terminänderung)
nicht vorhanden
""",
                {
                    "Segment Gruppe": "",
                    "Segment": "",
                    "Datenelement": "",
                    "Codes und Qualifier": "",
                    "Beschreibung": "",
                    "77777": "",
                    "88888": "",
                    "99999": "",
                    "Bedingung": """[12] Wenn SG4 DTM+471 (Ende zum nächstmöglichem Termin) nicht vorhanden  \n[13] Wenn SG4 STS+E01++Z01 (Status der Antwort: Zustimmung mit Terminänderung) nicht vorhanden """,
                },
                id="First Bedingung in UTILMD AHB WiM",
            ),
        ],
    )
    def test_parse_bedingung_cell(self, text_content, expected_df_row):

        # insert text
        self.test_cell.text = text_content

        # Initial two dataframes ...
        df = pd.DataFrame(columns=expected_df_row.keys(), dtype="str")
        expected_df = pd.DataFrame(columns=expected_df_row.keys(), dtype="str")
        row_index = 0
        # ... with a row full of emtpy strings
        initial_dataframe_row = (len(df.columns)) * [""]
        df.loc[row_index] = initial_dataframe_row
        expected_df.loc[row_index] = initial_dataframe_row

        parse_bedingung_cell(
            bedingung_cell=self.test_cell,
            dataframe=df,
            row_index=row_index,
        )

        expected_df.loc[row_index] = expected_df_row

        assert expected_df.equals(df)


@dataclass
class _Paragraph:
    text: str
    tabstop_positions: List[int]
    left_indent_position: int
    is_bold: bool = False


class TestWriteFunctions:

    # this left indent and tabstop positions are equal to
    # the left indent and tabstop positions of the indicator paragraph
    middle_cell_left_indent_position_of_indicator_paragraph = 36830
    middle_cell_tabstop_positions_of_indicator_paragraph = [436245, 1962785, 2578735, 3192780]

    edifact_struktur_cell_left_indent_position_of_indicator_paragraph = 364490
    edifact_struktur_cell_left_indent_position_of_segmentgroup_cells = 36830
    edifact_struktur_cell_tabstop_positions = [364490, 692150]

    def _prepare_docx_table_row(self, row_cells, expected_df_rows):
        # create table test cell
        # it contains per default an empty paragraph
        test_document = docx.Document()
        test_table = test_document.add_table(rows=1, cols=3)

        self.edifact_struktur_cell = test_table.add_row().cells[0]
        self.middle_cell = test_table.add_row().cells[1]
        self.bedingung_cell = test_table.add_row().cells[2]

        # prepare edifact struktur cell
        current_paragraph = self.edifact_struktur_cell.paragraphs[0]

        current_paragraph.text = row_cells["edifact_struktur_cell"][0].text
        current_paragraph.paragraph_format.left_indent = row_cells["edifact_struktur_cell"][0].left_indent_position

        if row_cells["edifact_struktur_cell"][0].tabstop_positions is not None:
            for tabstop_position in row_cells["edifact_struktur_cell"][0].tabstop_positions:
                self.edifact_struktur_cell.paragraphs[0].paragraph_format.tab_stops.add_tab_stop(tabstop_position)

        for _paragraph, i in zip(
            row_cells["edifact_struktur_cell"][1:], range(1, len(row_cells["edifact_struktur_cell"]))
        ):
            self.edifact_struktur_cell.add_paragraph()
            current_paragraph = self.edifact_struktur_cell.paragraphs[i]
            current_paragraph.text = _paragraph.text

            current_paragraph.paragraph_format.left_indent = _paragraph.left_indent_position
            current_tab_stops = current_paragraph.paragraph_format.tab_stops

            if _paragraph.tabstop_positions is not None:
                for tabstop_position in _paragraph.tabstop_positions:
                    current_tab_stops.add_tab_stop(tabstop_position)

        # prepare middle cell
        current_paragraph = self.middle_cell.paragraphs[0]

        current_paragraph.text = row_cells["middle_cell"][0].text
        current_paragraph.runs[0].bold = row_cells["middle_cell"][0].is_bold
        current_paragraph.paragraph_format.left_indent = row_cells["middle_cell"][0].left_indent_position

        if row_cells["middle_cell"][0].tabstop_positions is not None:
            for tabstop_position in row_cells["middle_cell"][0].tabstop_positions:
                self.middle_cell.paragraphs[0].paragraph_format.tab_stops.add_tab_stop(tabstop_position)

        for _paragraph, i in zip(row_cells["middle_cell"][1:], range(1, len(row_cells["middle_cell"]))):
            self.middle_cell.add_paragraph()
            current_paragraph = self.middle_cell.paragraphs[i]
            current_paragraph.text = _paragraph.text
            current_paragraph.runs[0].bold = row_cells["middle_cell"][i].is_bold

            current_paragraph.paragraph_format.left_indent = _paragraph.left_indent_position
            current_tab_stops = current_paragraph.paragraph_format.tab_stops

            if _paragraph.tabstop_positions is not None:
                for tabstop_position in _paragraph.tabstop_positions:
                    current_tab_stops.add_tab_stop(tabstop_position)

        # prepare bedingung cell
        current_paragraph = self.bedingung_cell.paragraphs[0]
        current_paragraph.text = row_cells["bedingung_cell"]

        # Initial two dataframes ...
        df = pd.DataFrame(columns=expected_df_rows[0].keys(), dtype="str")
        expected_df = pd.DataFrame(columns=expected_df_rows[0].keys(), dtype="str")
        row_index = 0
        # ... with a row full of emtpy strings
        initial_dataframe_row = (len(df.columns)) * [""]
        df.loc[row_index] = initial_dataframe_row
        expected_df.loc[row_index] = initial_dataframe_row

        return row_index, df, expected_df

    @pytest.mark.parametrize(
        "row_cells, expected_df_rows",
        [
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="Nachrichten-Kopfsegment",
                            tabstop_positions=None,
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                        )
                    ],
                    "middle_cell": [_Paragraph(text="", tabstop_positions=None, left_indent_position=None)],
                    "bedingung_cell": "",
                },
                [
                    {
                        "Segment Gruppe": "Nachrichten-Kopfsegment",
                        "Segment": "",
                        "Datenelement": "",
                        "Codes und Qualifier": "",
                        "Beschreibung": "",
                        "77777": "",
                        "88888": "",
                        "99999": "",
                        "Bedingung": "",
                    }
                ],
                id="Segmentname: Nachrichten-Kopfsegment",
            ),
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="Ende zum",
                            tabstop_positions=None,
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                        ),
                        _Paragraph(
                            text="(nächstmöglichem Termin)",
                            tabstop_positions=None,
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                        ),
                    ],
                    "middle_cell": [_Paragraph(text="", tabstop_positions=None, left_indent_position=None)],
                    "bedingung_cell": "",
                },
                [
                    {
                        "Segment Gruppe": "Ende zum (nächstmöglichem Termin)",
                        "Segment": "",
                        "Datenelement": "",
                        "Codes und Qualifier": "",
                        "Beschreibung": "",
                        "77777": "",
                        "88888": "",
                        "99999": "",
                        "Bedingung": "",
                    }
                ],
                id="Ende zum ...",
            ),
        ],
    )
    def test_write_segment_name_to_dataframe(
        self,
        row_cells,
        expected_df_rows,
    ):
        row_index, df, expected_df = self._prepare_docx_table_row(row_cells, expected_df_rows)

        write_segment_name_to_dataframe(
            dataframe=df,
            row_index=row_index,
            edifact_struktur_cell=self.edifact_struktur_cell,
            edifact_struktur_cell_left_indent_position=self.edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
            middle_cell=self.middle_cell,
            middle_cell_left_indent_position=self.middle_cell_left_indent_position_of_indicator_paragraph,
            tabstop_positions=self.middle_cell_tabstop_positions_of_indicator_paragraph,
            bedingung_cell=self.bedingung_cell,
        )

        expected_df.loc[row_index] = expected_df_rows[0]

        assert expected_df.equals(df)

    @pytest.mark.parametrize(
        "row_cells, expected_df_rows",
        [
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="SG2",
                            tabstop_positions=None,
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                        )
                    ],
                    "middle_cell": [
                        _Paragraph(
                            text="\tMuss\tMuss\tMuss",
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph[1:],
                            left_indent_position=None,
                        )
                    ],
                    "bedingung_cell": "",
                },
                [
                    {
                        "Segment Gruppe": "SG2",
                        "Segment": "",
                        "Datenelement": "",
                        "Codes und Qualifier": "",
                        "Beschreibung": "",
                        "77777": "Muss",
                        "88888": "Muss",
                        "99999": "Muss",
                        "Bedingung": "",
                    }
                ],
                id="SG2",
            ),
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="SG6",
                            tabstop_positions=None,
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                        )
                    ],
                    "middle_cell": [
                        _Paragraph(
                            text="\tMuss\tMuss",
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph[2:],
                            left_indent_position=None,
                        )
                    ],
                    "bedingung_cell": "",
                },
                [
                    {
                        "Segment Gruppe": "SG6",
                        "Segment": "",
                        "Datenelement": "",
                        "Codes und Qualifier": "",
                        "Beschreibung": "",
                        "77777": "",
                        "88888": "Muss",
                        "99999": "Muss",
                        "Bedingung": "",
                    }
                ],
                id="SG6 Referenz Vorgangsnummer",
            ),
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="SG8",
                            tabstop_positions=None,
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                        )
                    ],
                    "middle_cell": [
                        _Paragraph(
                            text="\tMuss [138]\tMuss",
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph[1:3],
                            left_indent_position=None,
                        )
                    ],
                    "bedingung_cell": """[138] Wenn SG5
LOC+172 (Meldepunkt)
nicht vorhanden""",
                },
                [
                    {
                        "Segment Gruppe": "SG8",
                        "Segment": "",
                        "Datenelement": "",
                        "Codes und Qualifier": "",
                        "Beschreibung": "",
                        "77777": "Muss [138]",
                        "88888": "Muss",
                        "99999": "",
                        "Bedingung": "[138] Wenn SG5 LOC+172 (Meldepunkt) nicht vorhanden",
                    }
                ],
                id="SG8",
            ),
        ],
    )
    def test_write_segmentgruppe_to_dataframe(
        self,
        row_cells,
        expected_df_rows,
    ):
        row_index, df, expected_df = self._prepare_docx_table_row(row_cells, expected_df_rows)

        write_segmentgruppe_to_dataframe(
            dataframe=df,
            row_index=row_index,
            edifact_struktur_cell=self.edifact_struktur_cell,
            edifact_struktur_cell_left_indent_position=self.edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
            middle_cell=self.middle_cell,
            middle_cell_left_indent_position=self.middle_cell_left_indent_position_of_indicator_paragraph,
            tabstop_positions=self.middle_cell_tabstop_positions_of_indicator_paragraph,
            bedingung_cell=self.bedingung_cell,
        )

        expected_df.loc[row_index] = expected_df_rows[0]

        assert expected_df.equals(df)

    @pytest.mark.parametrize(
        "row_cells, expected_df_rows",
        [
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="UNH",
                            tabstop_positions=None,
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
                        )
                    ],
                    "middle_cell": [
                        _Paragraph(
                            text="\tMuss\tMuss\tMuss",
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph[1:],
                            left_indent_position=None,
                        )
                    ],
                    "bedingung_cell": "",
                },
                [
                    {
                        "Segment Gruppe": "",
                        "Segment": "UNH",
                        "Datenelement": "",
                        "Codes und Qualifier": "",
                        "Beschreibung": "",
                        "77777": "Muss",
                        "88888": "Muss",
                        "99999": "Muss",
                        "Bedingung": "",
                    }
                ],
                id="UNH",
            ),
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="SG2\tNAD",
                            tabstop_positions=edifact_struktur_cell_tabstop_positions[:1],
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                        )
                    ],
                    "middle_cell": [
                        _Paragraph(
                            text="\tMuss\tMuss\tMuss",
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph[1:],
                            left_indent_position=None,
                        )
                    ],
                    "bedingung_cell": "",
                },
                [
                    {
                        "Segment Gruppe": "SG2",
                        "Segment": "NAD",
                        "Datenelement": "",
                        "Codes und Qualifier": "",
                        "Beschreibung": "",
                        "77777": "Muss",
                        "88888": "Muss",
                        "99999": "Muss",
                        "Bedingung": "",
                    }
                ],
                id="Segmentgroup, Segment and Bedingung",
            ),
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="SG4\tSTS",
                            tabstop_positions=edifact_struktur_cell_tabstop_positions[:1],
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                        )
                    ],
                    "middle_cell": [
                        _Paragraph(
                            text="\tMuss [249]\tMuss [249]",
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph[2:],
                            left_indent_position=None,
                        )
                    ],
                    "bedingung_cell": """[249] Innerhalb eines
SG4 IDE müssen alle
DE1131 der SG4
STS+E01 den
identischen Wert
enthalten""",
                },
                [
                    {
                        "Segment Gruppe": "SG4",
                        "Segment": "STS",
                        "Datenelement": "",
                        "Codes und Qualifier": "",
                        "Beschreibung": "",
                        "77777": "",
                        "88888": "Muss [249]",
                        "99999": "Muss [249]",
                        "Bedingung": "[249] Innerhalb eines SG4 IDE müssen alle DE1131 der SG4 STS+E01 den identischen Wert enthalten",
                    }
                ],
                id="Segmentgroup and Segment",
            ),
        ],
    )
    def test_write_segment_to_dataframe(
        self,
        row_cells,
        expected_df_rows,
    ):
        row_index, df, expected_df = self._prepare_docx_table_row(row_cells, expected_df_rows)

        write_segment_to_dataframe(
            dataframe=df,
            row_index=row_index,
            edifact_struktur_cell=self.edifact_struktur_cell,
            edifact_struktur_cell_left_indent_position=self.edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
            middle_cell=self.middle_cell,
            middle_cell_left_indent_position=self.middle_cell_left_indent_position_of_indicator_paragraph,
            tabstop_positions=self.middle_cell_tabstop_positions_of_indicator_paragraph,
            bedingung_cell=self.bedingung_cell,
        )

        expected_df.loc[row_index] = expected_df_rows[0]

        assert expected_df.equals(df)

    @pytest.mark.parametrize(
        "row_cells, expected_df_rows",
        [
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="UNH\t0062",
                            tabstop_positions=edifact_struktur_cell_tabstop_positions[1:],
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
                        )
                    ],
                    "middle_cell": [
                        _Paragraph(
                            text="Nachrichten-Referenznummer\tX\tX\tX",
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph[1:],
                            left_indent_position=middle_cell_left_indent_position_of_indicator_paragraph,
                        )
                    ],
                    "bedingung_cell": "",
                },
                [
                    {
                        "Segment Gruppe": "",
                        "Segment": "UNH",
                        "Datenelement": "0062",
                        "Codes und Qualifier": "Nachrichten-Referenznummer",
                        "Beschreibung": "",
                        "77777": "X",
                        "88888": "X",
                        "99999": "X",
                        "Bedingung": "",
                    }
                ],
                id="UNH\t0062",
            ),
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="UNH\t0065",
                            tabstop_positions=edifact_struktur_cell_tabstop_positions[:1],
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
                        )
                    ],
                    "middle_cell": [
                        _Paragraph(
                            text="UTILM\tNetzanschluss-\tX\tX\tX",
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph,
                            left_indent_position=middle_cell_left_indent_position_of_indicator_paragraph,
                        ),
                        _Paragraph(
                            text="D\tStammdaten",
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph[:1],
                            left_indent_position=middle_cell_left_indent_position_of_indicator_paragraph,
                        ),
                    ],
                    "bedingung_cell": "",
                },
                [
                    {
                        "Segment Gruppe": "",
                        "Segment": "UNH",
                        "Datenelement": "0065",
                        "Codes und Qualifier": "UTILMD",
                        "Beschreibung": "Netzanschluss-Stammdaten",
                        "77777": "X",
                        "88888": "X",
                        "99999": "X",
                        "Bedingung": "",
                    }
                ],
                id="UNH\t0065",
            ),
        ],
    )
    def test_single_line_write_dataelement_to_dataframe(
        self,
        row_cells,
        expected_df_rows,
    ):

        row_index, df, expected_df = self._prepare_docx_table_row(row_cells, expected_df_rows)

        write_dataelement_to_dataframe(
            dataframe=df,
            row_index=row_index,
            edifact_struktur_cell=self.edifact_struktur_cell,
            edifact_struktur_cell_left_indent_position=self.edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
            middle_cell=self.middle_cell,
            middle_cell_left_indent_position=self.middle_cell_left_indent_position_of_indicator_paragraph,
            tabstop_positions=self.middle_cell_tabstop_positions_of_indicator_paragraph,
            bedingung_cell=self.bedingung_cell,
        )

        expected_df.loc[row_index] = expected_df_rows[0]

        assert expected_df.equals(df)

    @pytest.mark.parametrize(
        "row_cells, expected_df_rows",
        [
            pytest.param(
                {
                    "edifact_struktur_cell": [
                        _Paragraph(
                            text="SG2\tNAD\t3055",
                            tabstop_positions=edifact_struktur_cell_tabstop_positions,
                            left_indent_position=edifact_struktur_cell_left_indent_position_of_segmentgroup_cells,
                        )
                    ],
                    "middle_cell": [
                        _Paragraph(
                            text="9\tGS1\tX\tX\tX",
                            is_bold=True,
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph,
                            left_indent_position=middle_cell_left_indent_position_of_indicator_paragraph,
                        ),
                        _Paragraph(
                            text="293\tDE, BDEW\tX\tX\tX",
                            is_bold=True,
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph,
                            left_indent_position=middle_cell_left_indent_position_of_indicator_paragraph,
                        ),
                        _Paragraph(
                            text="(Bundesverband der",
                            is_bold=False,
                            tabstop_positions=None,
                            left_indent_position=middle_cell_tabstop_positions_of_indicator_paragraph[0],
                        ),
                        _Paragraph(
                            text="Energie- und",
                            is_bold=False,
                            tabstop_positions=None,
                            left_indent_position=middle_cell_tabstop_positions_of_indicator_paragraph[0],
                        ),
                        _Paragraph(
                            text="Wasserwirtschaft e.V.)",
                            is_bold=False,
                            tabstop_positions=None,
                            left_indent_position=middle_cell_tabstop_positions_of_indicator_paragraph[0],
                        ),
                        _Paragraph(
                            text="332\tDE, DVGW Service &\tX\tX\tX",
                            is_bold=True,
                            tabstop_positions=middle_cell_tabstop_positions_of_indicator_paragraph,
                            left_indent_position=middle_cell_left_indent_position_of_indicator_paragraph,
                        ),
                        _Paragraph(
                            text="Consult GmbH",
                            is_bold=False,
                            tabstop_positions=None,
                            left_indent_position=middle_cell_tabstop_positions_of_indicator_paragraph[0],
                        ),
                    ],
                    "bedingung_cell": "",
                },
                [
                    {
                        "Segment Gruppe": "SG2",
                        "Segment": "NAD",
                        "Datenelement": "3055",
                        "Codes und Qualifier": "9",
                        "Beschreibung": "GS1",
                        "77777": "X",
                        "88888": "X",
                        "99999": "X",
                        "Bedingung": "",
                    },
                    {
                        "Segment Gruppe": "SG2",
                        "Segment": "NAD",
                        "Datenelement": "3055",
                        "Codes und Qualifier": "293",
                        "Beschreibung": "DE, BDEW (Bundesverband der Energie- und Wasserwirtschaft e.V.)",
                        "77777": "X",
                        "88888": "X",
                        "99999": "X",
                        "Bedingung": "",
                    },
                    {
                        "Segment Gruppe": "SG2",
                        "Segment": "NAD",
                        "Datenelement": "3055",
                        "Codes und Qualifier": "332",
                        "Beschreibung": "DE, DVGW Service & Consult GmbH",
                        "77777": "X",
                        "88888": "X",
                        "99999": "X",
                        "Bedingung": "",
                    },
                ],
                id="SG2\tNAD\t3055",
            ),
        ],
    )
    def test_multi_line_write_dataelement_to_dataframe(
        self,
        row_cells,
        expected_df_rows,
    ):
        row_index, df, expected_df = self._prepare_docx_table_row(row_cells, expected_df_rows)

        write_dataelement_to_dataframe(
            dataframe=df,
            row_index=row_index,
            edifact_struktur_cell=self.edifact_struktur_cell,
            edifact_struktur_cell_left_indent_position=self.edifact_struktur_cell_left_indent_position_of_indicator_paragraph,
            middle_cell=self.middle_cell,
            middle_cell_left_indent_position=self.middle_cell_left_indent_position_of_indicator_paragraph,
            tabstop_positions=self.middle_cell_tabstop_positions_of_indicator_paragraph,
            bedingung_cell=self.bedingung_cell,
        )

        for expected_df_row in expected_df_rows:

            expected_df.loc[row_index] = expected_df_row
            row_index += 1

        assert expected_df.equals(df)
