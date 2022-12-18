import docx
import pytest
from docx.shared import Inches, Twips
import pandas as pd

from kohlrahbi.cells import BodyCell


class TestBodyCell:
    """
    This test class contains all tests of the BodyCell class
    """

    @pytest.mark.parametrize(
        ["cell_text", "expected_dataframe"],
        [
            pytest.param(
                "Nachrichten-Referenznummer\tX\tX\tX",
                pd.DataFrame(
                    {
                        "Segment Gruppe": [""],
                        "Segment": [""],
                        "Datenelement": [""],
                        "Codes und Qualifier": ["Nachrichten-Referenznummer"],
                        "Beschreibung": [""],
                        "11016": ["X"],
                        "11017": ["X"],
                        "11018": ["X"],
                    }
                ),
                id="Nachrichten-Referenznummer",
            ),
        ],
    )
    def test_body_cell_parse(self, cell_text: str, expected_dataframe: pd.DataFrame):

        indicator_tabstop_positions: list[int] = [440055, 1960880, 2580640, 3191510]

        indicator_tabstop_positions_in_twips = []

        for position in indicator_tabstop_positions:
            indicator_tabstop_positions_in_twips.append(Twips(position / 635))

        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)
        body_cell = table.rows[0].cells
        body_cell[0].text = "Nachrichten-Referenznummer\tX\tX\tX"

        for tabstop_position in indicator_tabstop_positions_in_twips[1:]:
            body_cell[0].paragraphs[0].paragraph_format.tab_stops.add_tab_stop(tabstop_position)

        left_indent_length = Twips(693)
        body_cell[0].paragraphs[0].paragraph_format.left_indent = left_indent_length

        bc: BodyCell = BodyCell(
            table_cell=table.row_cells(0)[0],
            left_indent_position=left_indent_length,
            indicator_tabstop_positions=indicator_tabstop_positions,
        )

        df: pd.DataFrame = pd.DataFrame(
            {
                "Segment Gruppe": [""],
                "Segment": [""],
                "Datenelement": [""],
                "Codes und Qualifier": [""],
                "Beschreibung": [""],
                "11016": [""],
                "11017": [""],
                "11018": [""],
            }
        )
        df = bc.parse(ahb_row_dataframe=df)

        assert df.equals(expected_dataframe)
