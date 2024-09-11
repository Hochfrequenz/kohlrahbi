import json
from pathlib import Path

import docx
import pytest
from docx import Document
from efoli import EdifactFormat, EdifactFormatVersion

from kohlrahbi.ahb import get_pruefi_to_file_mapping
from kohlrahbi.ahbtable.ahbcondtions import AhbConditions
from kohlrahbi.ahbtable.ahbpackagetable import AhbPackageTable
from kohlrahbi.conditions import find_all_files_from_all_pruefis
from kohlrahbi.read_functions import get_all_conditions_from_doc, is_item_package_heading
from unittests import test_formats


def create_heading_paragraph(text, style):
    paragraph = Document().add_paragraph(text)
    paragraph.style = style
    return paragraph


@pytest.mark.snapshot
class TestReadFunctions:
    @pytest.mark.parametrize(
        "item, style_name, expected_boolean",
        [
            pytest.param(
                create_heading_paragraph("Übersicht der Pakete in der UTILMD", "Heading 1"),
                "Heading 1",
                True,
                id="heading1_valid",
            ),
            pytest.param(
                create_heading_paragraph("Übersicht der Pakete in der UTILMD", "Heading 2"),
                "Heading 2",
                True,
                id="heading2_valid",
            ),
            pytest.param(
                create_heading_paragraph("Übersicht der Pakete in der UTILMD", "Heading 3"),
                "Heading 3",
                False,
                id="invalid_style",
            ),
            pytest.param(
                create_heading_paragraph("Übersicht der Pakete in der IFSTA", "Heading 2"),
                "Heading 2",
                False,
                id="invalid_edifact_format",
            ),
            pytest.param(
                Document().add_table(rows=1, cols=1),
                "Heading 2",
                False,
                id="table_invalid",
            ),
        ],
    )
    def test_is_package_item_heading(self, item, style_name, expected_boolean):
        assert is_item_package_heading(item, style_name, EdifactFormat.UTILMD) == expected_boolean

    @pytest.mark.parametrize(
        "edifact_format",
        [
            pytest.param(
                format,
                id=str(format),
            )
            for format in test_formats
        ],
    )
    def test_get_all_conditions_from_doc(self, edifact_format: EdifactFormat, tmp_path, snapshot):

        edi_repo_path = Path(str(Path(__file__).parents[1] / "edi_energy_mirror"))
        pruefi_to_file_mapping = get_pruefi_to_file_mapping(edi_repo_path, EdifactFormatVersion.FV2404)

        collected_conditions: AhbConditions = AhbConditions()
        collected_packages: AhbPackageTable = AhbPackageTable()
        all_format_files = find_all_files_from_all_pruefis(pruefi_to_file_mapping)

        for file in all_format_files.get(edifact_format):
            # type: ignore[call-arg, arg-type]
            doc = docx.Document(edi_repo_path / "edi_energy_de" / "FV2404" / Path(file))
            assert doc
            packages, cond_table = get_all_conditions_from_doc(doc, edifact_format)
            if packages.table is not None:
                collected_conditions.include_condition_dict(packages.provide_conditions(edifact_format))
            collected_conditions.include_condition_dict(cond_table.conditions_dict)
            collected_packages.include_package_dict(packages.package_dict)
        assert collected_conditions == snapshot
        assert collected_packages == snapshot
