import docx  # type:ignore[import]
import pytest  # type:ignore[import]
from docx.shared import RGBColor  # type:ignore[import]

from ahbextractor.helper.check_row_type import RowType, define_row_type


class TestCheckRowType:

    # this is the touchstone position for the left indent of the edifact struktur cells
    # example rows with this left indent position:
    # |      UNH       |
    # |      UNH\t0062 |
    segment_left_indent_position = 364490

    # example rows with this left indent position:
    # | SG2            |
    # | SG2\tNAD       |
    # | SG2\tNAD\t3025 |
    segmentgruppe_left_indent_position = 36830

    @pytest.mark.parametrize(
        "text, left_indent_position, font_color, expected",
        [
            pytest.param(
                "EDIFACT Struktur", segmentgruppe_left_indent_position, RGBColor(0, 0, 0), RowType.HEADER, id="HEADER"
            ),
            pytest.param(
                "Nachrichten-Kopfsegment",
                segmentgruppe_left_indent_position,
                RGBColor(128, 128, 128),
                RowType.SEGMENTNAME,
                id="SEGMENTNAME",
            ),
            pytest.param(
                "SG2", segmentgruppe_left_indent_position, RGBColor(0, 0, 0), RowType.SEGMENTGRUPPE, id="SEGMENTGRUPPE"
            ),
            pytest.param(
                "UNH",
                segment_left_indent_position,
                RGBColor(0, 0, 0),
                RowType.SEGMENT,
                id="SEGMENT w/o Segmentgruppe",
            ),
            pytest.param(
                "SG2\tNAD",
                segmentgruppe_left_indent_position,
                RGBColor(0, 0, 0),
                RowType.SEGMENT,
                id="SEGMENT w Segementgruppe",
            ),
            pytest.param(
                "UNH\t0062",
                segment_left_indent_position,
                RGBColor(0, 0, 0),
                RowType.DATENELEMENT,
                id="DATENELEMENT w/o Segmentgruppe",
            ),
            pytest.param(
                "SG2\tNAD\t3035",
                segmentgruppe_left_indent_position,
                RGBColor(0, 0, 0),
                RowType.DATENELEMENT,
                id="DATENELEMENT w Segementgruppe",
            ),
            pytest.param("", 635, RGBColor(0, 0, 0), RowType.EMPTY, id="EMPTY"),
        ],  # TODO find left indent for empty cell
    )
    def test_define_row_type(self, text: str, left_indent_position: int, font_color: RGBColor, expected: RowType):
        """Test if all defined row types are identified correctly.

        Args:
            text (str): Text content of the test cell
            left_indent_position (int): Position of the left indent in arbritrary units
            font_color (RGBColor): A class from docx to define colors
            expected (RowType): The expected RowType for each test case
        """
        # ! Attention: It seems that you can set the left indent only to discret numbers!
        # edifact_struktur_left_indent_position = 1270

        # create table test cell, it contains per default an empty paragraph
        test_document = docx.Document()
        test_table = test_document.add_table(rows=1, cols=1)
        test_cell = test_table.add_row().cells[0]

        # insert text
        test_cell.text = text

        # set left indent positon
        test_cell.paragraphs[0].paragraph_format.left_indent = left_indent_position
        # set font color
        test_cell.paragraphs[0].runs[0].font.color.rgb = font_color

        result = define_row_type(
            edifact_struktur_cell=test_cell, left_indent_position=self.segment_left_indent_position
        )
        assert result == expected
