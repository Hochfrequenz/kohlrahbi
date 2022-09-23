"""
Collection of functions to write the extracted infos from the AHB tables into a DataFrame.
"""

import re
from typing import List

import pandas as pd
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from ahbextractor.helper.check_row_type import RowType
from ahbextractor.helper.elixir import Elixir


def parse_paragraph_in_edifact_struktur_column_to_dataframe(
    paragraphs: List[Paragraph],
    dataframe: pd.DataFrame,
    row_index: int,
    edifact_struktur_cell_left_indent_position: int,
):
    """Parses a paragraph in the edifact struktur column and puts the information into the appropriate columns

    Args:
        paragraphs (Paragraph): Current paragraphs in the edifact struktur cell
        dataframe (pd.DataFrame): Contains all infos
        row_index (int): Current index of the DataFrame
        edifact_struktur_cell_left_indent_position (int): Position of the left indent from the indicator edifact
            struktur cell
    """
    joined_text = " ".join(p.text for p in paragraphs)
    splitted_text_at_tabs = joined_text.split("\t")
    tab_count = joined_text.count("\t")

    # Check if the line starts on the far left
    if paragraphs[0].paragraph_format.left_indent != edifact_struktur_cell_left_indent_position:

        if tab_count == 2:
            dataframe.at[row_index, "Segment Gruppe"] = splitted_text_at_tabs[0]
            dataframe.at[row_index, "Segment"] = splitted_text_at_tabs[1]
            dataframe.at[row_index, "Datenelement"] = splitted_text_at_tabs[2]
        elif tab_count == 1:
            dataframe.at[row_index, "Segment Gruppe"] = splitted_text_at_tabs[0]
            dataframe.at[row_index, "Segment"] = splitted_text_at_tabs[1]
        elif tab_count == 0 and joined_text != "":
            if paragraphs[0].runs[0].bold:
                # Segmentgruppe: SG8
                dataframe.at[row_index, "Segment Gruppe"] = splitted_text_at_tabs[0]
            else:
                # Segmentname: Referenzen auf die ID der\nTranche
                _sg_text = dataframe.at[row_index, "Segment Gruppe"]
                if _sg_text == "":
                    # Referenzen auf die ID der
                    dataframe.at[row_index, "Segment Gruppe"] = splitted_text_at_tabs[0]
                else:
                    # Tranche
                    dataframe.at[row_index, "Segment Gruppe"] += " " + splitted_text_at_tabs[0]

    # Now the text should start in middle of the EDIFACT Struktur column
    else:

        if tab_count == 1:
            # Example: "UNH\t0062"
            dataframe.at[row_index, "Segment"] = splitted_text_at_tabs[0]
            dataframe.at[row_index, "Datenelement"] = splitted_text_at_tabs[1]

        elif tab_count == 0:
            # Example: "UNH"
            dataframe.at[row_index, "Segment"] = splitted_text_at_tabs[0]


def parse_paragraph_in_middle_column_to_dataframe(
    paragraph: Paragraph,
    dataframe: pd.DataFrame,
    row_index: int,
    left_indent_position: int,
    tabstop_positions: List[int],
):
    """Parses a paragraph in the middle column and puts the information into the appropriate columns

    Args:
        paragraph (Paragraph): Current paragraph in the edifact struktur cell
        dataframe (pd.DataFrame): Contains all infos
        row_index (int): Current index of the DataFrame
        left_indent_position (int): Position of the left indent from the indicator middle cell
        tabstop_positions (List[int]): All tabstop positions of the indicator middle cell
    """

    splitted_text_at_tabs = paragraph.text.split("\t")

    # Qualifier / Code
    # left_indent_position is characteristic for Datenelemente
    if paragraph.paragraph_format.left_indent == left_indent_position:
        dataframe.at[row_index, "Codes und Qualifier"] += splitted_text_at_tabs.pop(0)
        column_indezes = list(range(4, 4 + len(tabstop_positions)))

    else:
        if splitted_text_at_tabs[0] == "":
            tabstop_positions = tabstop_positions[1:]
            del splitted_text_at_tabs[0]

        column_indezes = list(range(5, 5 + len(tabstop_positions)))

    # pylint: disable=protected-access
    tab_stops = paragraph.paragraph_format.tab_stops._pPr.tabs

    if tab_stops is not None:
        for tabstop in tab_stops:
            for tabstop_position, column_index in zip(tabstop_positions, column_indezes):
                if tabstop.pos == tabstop_position:
                    dataframe.iat[row_index, column_index] += splitted_text_at_tabs.pop(0)
    elif tab_stops is None and splitted_text_at_tabs:
        # in splitted_text_at_tabs list must be an entry
        dataframe.at[row_index, "Beschreibung"] += splitted_text_at_tabs.pop(0)
    elif tab_stops is None:
        pass
    # Could not figure out a scenario where this error could be raised.
    # else:
    #     raise NotImplementedError(f"Could not parse paragraph in middle cell with {paragraph.text}")


def parse_bedingung_cell(bedingung_cell: _Cell, dataframe: pd.DataFrame, row_index: int):
    """Parses a cell in the Bedingung column and puts the information into the in the appropriate column

    Args:
        bedingung_cell (_Cell): Cell from the Bedingung column
        dataframe (pd.DataFrame): Saves all infos
        row_index (int): Current index of the DataFrame
    """

    bedingung = bedingung_cell.text.replace("\n", " ")
    matches = re.findall(r"\[\d*\]", bedingung)
    for match in matches[1:]:
        index = bedingung.find(match)
        bedingung = bedingung[:index] + "\n" + bedingung[index:]

    dataframe.at[row_index, "Bedingung"] += bedingung


# pylint: disable=too-many-arguments
def write_segment_name_to_dataframe(
    elixir: Elixir,
    edifact_struktur_cell: _Cell,
    middle_cell: _Cell,
    bedingung_cell: _Cell,
):
    """Writes all infos from a segment name row into a DataFrame.

    Args:
        elixir (Elixir):
        edifact_struktur_cell (_Cell): Cell from the edifact struktur column
        middle_cell (_Cell): Cell from the middle column
        bedingung_cell (_Cell): Cell from the Bedingung column
    """

    # EDIFACT STRUKTUR COLUMN
    for paragraph in edifact_struktur_cell.paragraphs:
        parse_paragraph_in_edifact_struktur_column_to_dataframe(
            paragraphs=[paragraph],
            dataframe=elixir.soul,
            row_index=elixir.current_df_row_index,
            edifact_struktur_cell_left_indent_position=elixir.edifact_struktur_left_indent_position,
        )

    # MIDDLE COLUMN
    # I do not expect to get a multiline Segmentgruppe,
    # but just in case we loop through all paragraphs
    for paragraph in middle_cell.paragraphs:
        parse_paragraph_in_middle_column_to_dataframe(
            paragraph=paragraph,
            dataframe=elixir.soul,
            row_index=elixir.current_df_row_index,
            left_indent_position=elixir.middle_cell_left_indent_position,
            tabstop_positions=elixir.tabstop_positions,
        )

    # BEDINGUNG COLUMN
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=elixir.soul, row_index=elixir.current_df_row_index)


# pylint: disable=too-many-arguments
def write_segmentgruppe_to_dataframe(
    elixir: Elixir,
    edifact_struktur_cell: _Cell,
    middle_cell: _Cell,
    bedingung_cell: _Cell,
):
    """Writes all infos from a segmentgruppe row into a DataFrame.

    Args:
        elixir (Elixir):
        edifact_struktur_cell (_Cell): Cell from the edifact struktur column
        middle_cell (_Cell): Cell from the middle column
        bedingung_cell (_Cell): Cell from the Bedingung column
    """

    # EDIFACT STRUKTUR COLUMN
    parse_paragraph_in_edifact_struktur_column_to_dataframe(
        # there might be 2 paragraphs in case of multi line headings, so we're handing over all the paragraphs
        paragraphs=edifact_struktur_cell.paragraphs,
        dataframe=elixir.soul,
        row_index=elixir.current_df_row_index,
        edifact_struktur_cell_left_indent_position=elixir.edifact_struktur_left_indent_position,
    )

    # MIDDLE COLUMN
    # I do not expect to a multiline Segementgruppe,
    # but just in case we loop through all paragraphs
    for paragraph in middle_cell.paragraphs:
        parse_paragraph_in_middle_column_to_dataframe(
            paragraph=paragraph,
            dataframe=elixir.soul,
            row_index=elixir.current_df_row_index,
            left_indent_position=elixir.middle_cell_left_indent_position,
            tabstop_positions=elixir.tabstop_positions,
        )

    # BEDINGUNG COLUMN
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=elixir.soul, row_index=elixir.current_df_row_index)


# pylint: disable=too-many-arguments
def write_segment_to_dataframe(
    elixir: Elixir,
    edifact_struktur_cell: _Cell,
    middle_cell: _Cell,
    bedingung_cell: _Cell,
):
    """Writes all infos from a segment row into a DataFrame.

    Args:
        elixir (Elixir):
        edifact_struktur_cell (_Cell): Cell from the edifact struktur column
        middle_cell (_Cell): Cell from the middle column
        bedingung_cell (_Cell): Cell from the Bedingung column
    """

    # EDIFACT STRUKTUR COLUMN
    parse_paragraph_in_edifact_struktur_column_to_dataframe(
        paragraphs=edifact_struktur_cell.paragraphs,
        dataframe=elixir.soul,
        row_index=elixir.current_df_row_index,
        edifact_struktur_cell_left_indent_position=elixir.edifact_struktur_left_indent_position,
    )

    # MIDDLE COLUMN
    for paragraph in middle_cell.paragraphs:
        parse_paragraph_in_middle_column_to_dataframe(
            paragraph=paragraph,
            dataframe=elixir.soul,
            row_index=elixir.current_df_row_index,
            left_indent_position=elixir.middle_cell_left_indent_position,
            tabstop_positions=elixir.tabstop_positions,
        )

    # BEDINGUNG COLUMN
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=elixir.soul, row_index=elixir.current_df_row_index)


def count_matching(condition, condition_argument, seq):
    """Returns the amount of items in seq that return true from condition"""
    return sum(condition(item, condition_argument) for item in seq)


def code_condition(paragraph: Paragraph, pruefi_tabstops: List[int]) -> bool:
    """Checks if the paragraph contains a Code by checking for bold style.

    Example for Codes: UTILMD, 11A, UN,


    Args:
        paragraph (Paragraph): Current paragraph
        pruefi_tabstops (List[int]): All tabstop positions of the indicator middle cell

    Returns:
        [bool]:
    """
    try:
        # pylint: disable=protected-access
        tabstop_positions = [tab_position.pos for tab_position in paragraph.paragraph_format.tab_stops._pPr.tabs]
    except TypeError:
        return False

    if paragraph.runs[0].bold is True and any(x in tabstop_positions for x in pruefi_tabstops):
        return True
    return False


def has_middle_cell_multiple_codes(paragraphs: List[Paragraph], pruefi_tabstops: List[int]) -> bool:
    """Checks if the paragraphs of a middle cell contains more than one Code.

    Args:
        paragraphs (List[Paragraph]): All paragraphs in the current middle cell
        pruefi_tabstops (List[int]): All tabstop positions of the indicator middle cell

    Returns:
        bool:
    """

    if count_matching(condition=code_condition, condition_argument=pruefi_tabstops, seq=paragraphs) > 1:
        return True
    return False


# pylint: disable=too-many-arguments
def write_dataelement_to_dataframe(
    elixir: Elixir,
    edifact_struktur_cell: _Cell,
    middle_cell: _Cell,
    bedingung_cell: _Cell,
):
    """Writes all infos from a dataelement row into a DataFrame.

    Args:
        elixir (Elixir):
        edifact_struktur_cell (_Cell): Cell from the edifact struktur column
        middle_cell (_Cell): Cell from the middle column
        bedingung_cell (_Cell): Cell from the Bedingung column
    """

    # EDIFACT STRUKTUR COLUMN
    parse_paragraph_in_edifact_struktur_column_to_dataframe(
        paragraphs=edifact_struktur_cell.paragraphs,
        dataframe=elixir.soul,
        row_index=elixir.current_df_row_index,
        edifact_struktur_cell_left_indent_position=elixir.edifact_struktur_left_indent_position,
    )

    # BEDINGUNG COLUMN
    # Bedingung have to be parsed before MIDDLE COLUMN
    # The cell with multiple codes counts up the row_index
    # This will cause then an IndexError for Bedingung
    parse_bedingung_cell(bedingung_cell=bedingung_cell, dataframe=elixir.soul, row_index=elixir.current_df_row_index)

    # MIDDLE COLUMN
    if not has_middle_cell_multiple_codes(
        paragraphs=middle_cell.paragraphs, pruefi_tabstops=elixir.tabstop_positions[1:]
    ):
        for paragraph in middle_cell.paragraphs:
            parse_paragraph_in_middle_column_to_dataframe(
                paragraph=paragraph,
                dataframe=elixir.soul,
                row_index=elixir.current_df_row_index,
                left_indent_position=elixir.middle_cell_left_indent_position,
                tabstop_positions=elixir.tabstop_positions,
            )
        elixir.current_df_row_index = elixir.current_df_row_index + 1

    else:
        # The middle cell contains multiple Codes

        # here we have to look into the next row to see, if we have to add a new datarow or
        # if we have to collect more information in the next row which we have to add to the current row

        create_new_dataframe_row_indicator_list: List = [
            paragraph.runs[0].bold is True for paragraph in middle_cell.paragraphs
        ]

        for paragraph, i in zip(middle_cell.paragraphs, range(len(create_new_dataframe_row_indicator_list))):

            # For reasons of good readability the EDIFACT Struktur information gets written again

            # EDIFACT STRUKTUR COLUMN

            if edifact_struktur_cell.paragraphs[0].text != "":
                parse_paragraph_in_edifact_struktur_column_to_dataframe(
                    paragraphs=edifact_struktur_cell.paragraphs,
                    dataframe=elixir.soul,
                    row_index=elixir.current_df_row_index,
                    edifact_struktur_cell_left_indent_position=elixir.edifact_struktur_left_indent_position,
                )
            else:
                elixir.soul.at[elixir.current_df_row_index, "Segment Gruppe"] = elixir.soul.loc[
                    elixir.current_df_row_index - 1, "Segment Gruppe"
                ]
                elixir.soul.at[elixir.current_df_row_index, "Segment"] = elixir.soul.loc[
                    elixir.current_df_row_index - 1, "Segment"
                ]
                elixir.soul.at[elixir.current_df_row_index, "Datenelement"] = elixir.soul.loc[
                    elixir.current_df_row_index - 1, "Datenelement"
                ]

            if paragraph.runs[0].bold:
                parse_paragraph_in_middle_column_to_dataframe(
                    paragraph=paragraph,
                    dataframe=elixir.soul,
                    row_index=elixir.current_df_row_index,
                    left_indent_position=elixir.middle_cell_left_indent_position,
                    tabstop_positions=elixir.tabstop_positions,
                )

            elif paragraph.paragraph_format.left_indent == elixir.tabstop_positions[0]:
                # multi line Beschreibung
                elixir.soul.at[elixir.current_df_row_index, "Beschreibung"] += " " + paragraph.text

            if len(create_new_dataframe_row_indicator_list) > i + 1:
                if create_new_dataframe_row_indicator_list[i + 1]:
                    elixir.current_df_row_index = elixir.current_df_row_index + 1
                    elixir.soul.loc[elixir.current_df_row_index] = (len(elixir.soul.columns)) * [""]
            else:
                elixir.current_df_row_index = elixir.current_df_row_index + 1

    return elixir.current_df_row_index


def write_new_row_in_dataframe(
    elixir: Elixir,
    row_type: RowType,
    table: Table,
    row: int,
    index_for_middle_column: int,
):
    """Writes the current row of the current table into the DataFrame depending on the type of the row

    Args:
        elixir (Elixir):
        row_type (RowType): Type of the current row
        table (Table): Current table
        row (int): Row of the current table
        index_for_middle_column (int): Index of the actual middle column
    """

    if row_type is RowType.HEADER:
        pass

    elif row_type is RowType.SEGMENTNAME:
        write_segment_name_to_dataframe(
            elixir=elixir,
            edifact_struktur_cell=table.row_cells(row)[0],
            middle_cell=table.row_cells(row)[index_for_middle_column],
            bedingung_cell=table.row_cells(row)[-1],
        )
        elixir.current_df_row_index = elixir.current_df_row_index + 1

    elif row_type is RowType.SEGMENTGRUPPE:
        write_segmentgruppe_to_dataframe(
            elixir=elixir,
            edifact_struktur_cell=table.row_cells(row)[0],
            middle_cell=table.row_cells(row)[index_for_middle_column],
            bedingung_cell=table.row_cells(row)[-1],
        )
        elixir.current_df_row_index = elixir.current_df_row_index + 1

    elif row_type is RowType.SEGMENT:
        write_segment_to_dataframe(
            elixir=elixir,
            edifact_struktur_cell=table.row_cells(row)[0],
            middle_cell=table.row_cells(row)[index_for_middle_column],
            bedingung_cell=table.row_cells(row)[-1],
        )
        elixir.current_df_row_index = elixir.current_df_row_index + 1

    elif row_type is RowType.DATENELEMENT:
        write_dataelement_to_dataframe(
            elixir=elixir,
            edifact_struktur_cell=table.row_cells(row)[0],
            middle_cell=table.row_cells(row)[index_for_middle_column],
            bedingung_cell=table.row_cells(row)[-1],
        )

    elif row_type is RowType.EMPTY:
        pass
